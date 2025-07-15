import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
import os
import sys
import json
from datetime import datetime
from docx import Document # For Word report
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class CleanItApp:
    CONFIG_FILE = os.path.join(os.path.expanduser('~'), '.cleanit_config.json')

    def __init__(self, root):
        self.root = root
        self.root.title("CleanIt")

        # --- Fullscreen/Maximized Window ---
        try:
            self.root.state('zoomed')
        except tk.TclError:
            pass

        self.root.resizable(True, True)

        # --- Colors ---
        self.bg_color = "#E6F7FF"
        self.fg_color = "#000080"

        self.root.configure(bg=self.bg_color)

        # --- Styles for ttk widgets ---
        self.style = ttk.Style()
        self.style.theme_use('clam')
        self.style.configure('TFrame', background=self.bg_color)
        self.style.configure('TLabel', background=self.bg_color, foreground=self.fg_color, font=('Helvetica', 10))
        self.style.configure('TButton', background='#A3D9F7', foreground=self.fg_color, font=('Helvetica', 10, 'bold'), padding=5)
        self.style.map('TButton',
                       background=[('active', '#8BCDF2')],
                       foreground=[('active', self.fg_color)])
        self.style.configure('TEntry', fieldbackground='white', foreground='black', font=('Helvetica', 10))
        self.style.configure('TCombobox', fieldbackground='white', foreground='black', font=('Helvetica', 10))
        self.style.map('TCombobox',
                       fieldbackground=[('readonly', 'white')],
                       selectbackground=[('readonly', 'white')],
                       selectforeground=[('readonly', 'black')])
        self.style.configure('TRadiobutton', background=self.bg_color, foreground=self.fg_color)
        self.style.configure('TCheckbutton', background=self.bg_color, foreground=self.fg_color)
        self.style.configure('TNotebook', background=self.bg_color, borderwidth=0)
        self.style.configure('TNotebook.Tab', background='#C2E5FF', foreground=self.fg_color, font=('Helvetica', 10, 'bold'))
        self.style.map('TNotebook.Tab',
                       background=[('selected', self.bg_color), ('active', '#A3D9F7')],
                       foreground=[('selected', self.fg_color)])

        # Treeview style for missing values and duplicate review
        self.style.configure("Treeview",
                             background="white",
                             foreground="black",
                             rowheight=25,
                             fieldbackground="white")
        self.style.map('Treeview',
                       background=[('selected', self.fg_color)],
                       foreground=[('selected', 'white')])
        self.style.configure("Treeview.Heading",
                             font=('Helvetica', 10, 'bold'),
                             background='#A3D9F7',
                             foreground=self.fg_color)

        # --- Variables ---
        self.input_file_path = tk.StringVar()
        self.output_folder_path = tk.StringVar()
        self.df = None # The DataFrame itself, now a class attribute

        # Output Format Variables
        self.output_file_format = tk.StringVar(value=".xlsx") # Default to XLSX
        self.generate_report = tk.BooleanVar(value=True) # Default to generate report

        # Date format options
        self._date_format_options = [
            ("MM/DD/YYYY (e.g., 01/15/2023)", "%m/%d/%Y"),
            ("Weekday, Month DD, YYYY (e.g., Sunday, January 15, 2023)", "%A, %B %d, %Y"),
            ("YYYY-MM-DD (e.g., 2023-01-15)", "%Y-%m-%d"),
            ("MM/DD (e.g., 01/15)", "%m/%d"),
            ("MM/DD/YY (e.g., 01/15/23)", "%m/%d/%y"),
            ("DD-Mon (e.g., 15-Jan)", "%d-%b"),
            ("DD-Mon-YY (e.g., 15-Jan-23)", "%d-%b-%y"),
            ("YY-Mon-DD (e.g., 23-Jan-15)", "%y-%b-%d"),
            ("Mon-DD (e.g., Jan-15)", "%b-%d"),
            ("Month-DD (e.g., January-15)", "%B-%d"),
            ("Month DD, YYYY (e.g., January 15, 2023)", "%B %d, %Y"),
            ("YYYY-MM-DD HH:MM:SS (24-hour, e.g., 2023-01-15 13:30:45)", "%Y-%m-%d %H:%M:%S"),
            ("MM/DD/YYYY HH:MM AM/PM (e.g., 01/15/2023 01:30 PM)", "%m/%d/%Y %I:%M %p")
        ]
        self.date_format_display_strings = [item[0] for item in self._date_format_options]
        self._date_format_map = {item[0]: item[1] for item in self._date_format_options}
        self.selected_date_format_display = tk.StringVar(value=self.date_format_display_strings[0])

        # Sorting variables
        self.sort_column = tk.StringVar()
        self.sort_order = tk.BooleanVar(value=True) # True for Ascending, False for Descending

        # Cleaning Options Checkbox Variables
        self.do_trim_whitespace = tk.BooleanVar(value=True)
        self.do_capitalize_strings = tk.BooleanVar(value=True)
        self.do_remove_duplicates = tk.BooleanVar(value=True)

        # Dynamic column checkboxes
        self.date_column_vars = {} # Stores {column_name: tk.BooleanVar} for date formatting
        self.duplicate_check_column_vars = {} # Stores {column_name: tk.BooleanVar} for duplicate detection

        # Missing values Treeview related
        self.missing_values_tree = None # Will be initialized in create_widgets
        
        # Duplicate Review related
        self.duplicate_review_tree = None # Will be initialized in create_widgets
        self.duplicate_rows_to_keep_vars = {} # {original_idx: tk.BooleanVar} for rows in self.duplicate_review_df
        self.duplicate_review_df = None # DataFrame containing only rows involved in duplicates

        # Processing Metrics (for report)
        self.processing_metrics = {}
        
        self._load_config() # Load paths on startup
        self.create_widgets()

    def _load_config(self):
        """Loads last used paths from a config file."""
        if os.path.exists(self.CONFIG_FILE):
            try:
                with open(self.CONFIG_FILE, 'r') as f:
                    config = json.load(f)
                    self.input_file_path.set(config.get('last_input_file_path', ''))
                    self.output_folder_path.set(config.get('last_output_folder', ''))
                    self.output_file_format.set(config.get('last_output_format', '.xlsx'))
                    self.generate_report.set(config.get('generate_report', True))
            except Exception as e:
                messagebox.showwarning("Config Load Error", f"Could not load configuration: {e}")

    def _save_config(self):
        """Saves current paths and settings to a config file."""
        config = {
            'last_input_file_path': self.input_file_path.get(),
            'last_output_folder': self.output_folder_path.get(),
            'last_output_format': self.output_file_format.get(),
            'generate_report': self.generate_report.get()
        }
        try:
            with open(self.CONFIG_FILE, 'w') as f:
                json.dump(config, f)
        except Exception as e:
            messagebox.showwarning("Config Save Error", f"Could not save configuration: {e}")

    def create_widgets(self):
        # --- Main Frame ---
        main_frame = ttk.Frame(self.root, padding="20 20 20 20")
        main_frame.pack(expand=True, fill='both')

        # --- Title ---
        title_label = ttk.Label(main_frame, text="CleanIt: Data Cleaning & Formatting",
                                font=('Helvetica', 16, 'bold'), foreground=self.fg_color)
        title_label.pack(pady=(0, 20), fill='x')

        # --- Notebook (Tabs) ---
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(expand=True, fill='both')

        # --- Tab 1: File & Output ---
        file_output_tab = ttk.Frame(self.notebook, padding="15 15 15 15")
        self.notebook.add(file_output_tab, text="1. File & Output")

        file_output_tab.grid_columnconfigure(1, weight=1)
        file_output_tab.grid_rowconfigure(0, weight=1)
        file_output_tab.grid_rowconfigure(1, weight=1)
        file_output_tab.grid_rowconfigure(2, weight=1) # For output format
        file_output_tab.grid_rowconfigure(3, weight=1) # For report option

        ttk.Label(file_output_tab, text="Input File:").grid(row=0, column=0, sticky='w', pady=5)
        self.input_entry = ttk.Entry(file_output_tab, textvariable=self.input_file_path, state='readonly')
        self.input_entry.grid(row=0, column=1, padx=5, pady=5, sticky='ew')
        ttk.Button(file_output_tab, text="Browse...", command=self.browse_input_file).grid(row=0, column=2, padx=5, pady=5)

        ttk.Label(file_output_tab, text="Output Folder:").grid(row=1, column=0, sticky='w', pady=5)
        self.output_entry = ttk.Entry(file_output_tab, textvariable=self.output_folder_path, state='readonly')
        self.output_entry.grid(row=1, column=1, padx=5, pady=5, sticky='ew')
        ttk.Button(file_output_tab, text="Browse...", command=self.browse_output_folder).grid(row=1, column=2, padx=5, pady=5)

        ttk.Label(file_output_tab, text="Output File Format:").grid(row=2, column=0, sticky='w', pady=5)
        output_format_frame = ttk.Frame(file_output_tab)
        output_format_frame.grid(row=2, column=1, columnspan=2, sticky='w', padx=5, pady=5)
        ttk.Radiobutton(output_format_frame, text=".xlsx (Excel)", variable=self.output_file_format, value=".xlsx").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(output_format_frame, text=".csv (CSV)", variable=self.output_file_format, value=".csv").pack(side=tk.LEFT, padx=5)

        ttk.Checkbutton(file_output_tab, text="Generate Processing Report (.docx)", variable=self.generate_report).grid(row=3, column=0, columnspan=3, sticky='w', pady=10)


        # --- Tab 2: Cleaning Options ---
        cleaning_options_tab = ttk.Frame(self.notebook, padding="15 15 15 15")
        self.notebook.add(cleaning_options_tab, text="2. Cleaning Options")
        cleaning_options_tab.grid_columnconfigure(0, weight=1)

        row_counter = 0

        # Whitespace Trimming
        ttk.Checkbutton(cleaning_options_tab, text="Trim Leading/Trailing Whitespace", variable=self.do_trim_whitespace).grid(row=row_counter, column=0, sticky='w', pady=(0,2))
        row_counter += 1
        ttk.Label(cleaning_options_tab, text="  - Removes any spaces at the beginning or end of text values in all columns.",
                  font=('Helvetica', 9, 'italic'), foreground='gray', wraplength=500, justify='left').grid(row=row_counter, column=0, sticky='w', padx=10, pady=(0,10))
        row_counter += 1

        # Capitalization
        ttk.Checkbutton(cleaning_options_tab, text="Capitalize Text Fields (Title Case)", variable=self.do_capitalize_strings).grid(row=row_counter, column=0, sticky='w', pady=(10,2))
        row_counter += 1
        ttk.Label(cleaning_options_tab, text="  - Converts the first letter of each word in text fields to uppercase (e.g., 'john doe' -> 'John Doe').",
                  font=('Helvetica', 9, 'italic'), foreground='gray', wraplength=500, justify='left').grid(row=row_counter, column=0, sticky='w', padx=10, pady=(0,10))
        row_counter += 1
        
        # Duplicate Handling
        ttk.Checkbutton(cleaning_options_tab, text="Remove Duplicate Rows", variable=self.do_remove_duplicates).grid(row=row_counter, column=0, sticky='w', pady=(10,2))
        row_counter += 1
        ttk.Label(cleaning_options_tab, text="  - Identifies and removes exact duplicate rows, keeping the first occurrence.",
                  font=('Helvetica', 9, 'italic'), foreground='gray', wraplength=500, justify='left').grid(row=row_counter, column=0, sticky='w', padx=10, pady=(0,5))
        row_counter += 1
        ttk.Label(cleaning_options_tab, text="Select columns to check for duplicates (all must match):",
                  font=('Helvetica', 10, 'bold'), foreground=self.fg_color).grid(row=row_counter, column=0, sticky='w', pady=(5,5))
        row_counter += 1

        # Frame for dynamic duplicate check column checkboxes with scrollbar
        self.duplicate_check_columns_outer_frame = ttk.Frame(cleaning_options_tab)
        self.duplicate_check_columns_outer_frame.grid(row=row_counter, column=0, sticky='nsew', padx=5, pady=5)
        self.duplicate_check_columns_outer_frame.grid_columnconfigure(0, weight=1)
        self.duplicate_check_columns_outer_frame.grid_rowconfigure(0, weight=1)

        self.duplicate_check_columns_canvas = tk.Canvas(self.duplicate_check_columns_outer_frame, background=self.bg_color, highlightthickness=0, height=100) # Fixed height
        self.duplicate_check_columns_canvas.grid(row=0, column=0, sticky='nsew')

        self.duplicate_check_columns_scrollbar = ttk.Scrollbar(self.duplicate_check_columns_outer_frame, orient="vertical", command=self.duplicate_check_columns_canvas.yview)
        self.duplicate_check_columns_scrollbar.grid(row=0, column=1, sticky='ns')

        self.duplicate_check_columns_canvas.configure(yscrollcommand=self.duplicate_check_columns_scrollbar.set)
        self.duplicate_check_columns_canvas.bind('<Configure>', lambda e: self.duplicate_check_columns_canvas.configure(scrollregion = self.duplicate_check_columns_canvas.bbox("all")))
        self.duplicate_check_columns_canvas.bind_all("<MouseWheel>", self._on_mousewheel_duplicate_check)

        self.duplicate_check_columns_frame = ttk.Frame(self.duplicate_check_columns_canvas, padding="5 5 5 5")
        self.duplicate_check_columns_canvas.create_window((0, 0), window=self.duplicate_check_columns_frame, anchor="nw")
        self.duplicate_check_columns_frame.bind("<Configure>", lambda e: self.duplicate_check_columns_canvas.configure(scrollregion=self.duplicate_check_columns_canvas.bbox("all")))

        ttk.Label(self.duplicate_check_columns_frame, text="Load a file to see columns...", background=self.bg_color).pack(pady=10)


        # --- Tab 3: Duplicate Review ---
        duplicate_review_tab = ttk.Frame(self.notebook, padding="15 15 15 15")
        self.notebook.add(duplicate_review_tab, text="3. Duplicate Review")
        duplicate_review_tab.grid_columnconfigure(0, weight=1)
        duplicate_review_tab.grid_rowconfigure(1, weight=1)

        ttk.Label(duplicate_review_tab, text="Duplicate Rows Found (Check to KEEP, Uncheck to DELETE):",
                  font=('Helvetica', 10, 'bold'), foreground=self.fg_color).grid(row=0, column=0, sticky='w', pady=(0,5))

        # Treeview for duplicate review
        duplicate_tree_frame = ttk.Frame(duplicate_review_tab)
        duplicate_tree_frame.grid(row=1, column=0, sticky='nsew', pady=5)
        duplicate_tree_frame.grid_rowconfigure(0, weight=1)
        duplicate_tree_frame.grid_columnconfigure(0, weight=1)

        self.duplicate_review_tree = ttk.Treeview(duplicate_tree_frame, show='headings')
        self.duplicate_review_tree.grid(row=0, column=0, sticky='nsew')

        vsb_dup = ttk.Scrollbar(duplicate_tree_frame, orient="vertical", command=self.duplicate_review_tree.yview)
        vsb_dup.grid(row=0, column=1, sticky='ns')
        hsb_dup = ttk.Scrollbar(duplicate_tree_frame, orient="horizontal", command=self.duplicate_review_tree.xview)
        hsb_dup.grid(row=1, column=0, sticky='ew')

        self.duplicate_review_tree.configure(yscrollcommand=vsb_dup.set, xscrollcommand=hsb_dup.set)
        self.duplicate_review_tree.bind("<Button-1>", self._on_duplicate_treeview_click) # Bind click for toggling checkbox

        self.apply_duplicates_button = ttk.Button(duplicate_review_tab, text="Apply Duplicate Changes & Continue", command=self._apply_duplicate_changes)
        self.apply_duplicates_button.grid(row=2, column=0, pady=10)
        self.apply_duplicates_button.config(state='disabled') # Disabled until duplicates are found


        # --- Tab 4: Date Formatting ---
        date_formatting_tab = ttk.Frame(self.notebook, padding="15 15 15 15")
        self.notebook.add(date_formatting_tab, text="4. Date Formatting")

        date_formatting_tab.grid_columnconfigure(1, weight=1)
        date_formatting_tab.grid_rowconfigure(2, weight=1)

        ttk.Label(date_formatting_tab, text="Choose Output Date Format:").grid(row=0, column=0, sticky='w', pady=5)
        self.date_format_combobox = ttk.Combobox(date_formatting_tab, textvariable=self.selected_date_format_display,
                                                 values=self.date_format_display_strings, state='readonly')
        self.date_format_combobox.grid(row=0, column=1, padx=5, pady=5, sticky='ew')
        if self.date_format_display_strings:
            self.selected_date_format_display.set(self.date_format_display_strings[0])

        ttk.Label(date_formatting_tab, text="Select Columns to Format as Dates:").grid(row=1, column=0, sticky='nw', pady=5)
        
        self.date_columns_outer_frame = ttk.Frame(date_formatting_tab)
        self.date_columns_outer_frame.grid(row=2, column=0, columnspan=2, padx=5, pady=5, sticky='nsew')
        self.date_columns_outer_frame.grid_columnconfigure(0, weight=1)
        self.date_columns_outer_frame.grid_rowconfigure(0, weight=1)

        self.date_columns_canvas = tk.Canvas(self.date_columns_outer_frame, background=self.bg_color, highlightthickness=0)
        self.date_columns_canvas.grid(row=0, column=0, sticky='nsew')

        self.date_columns_scrollbar = ttk.Scrollbar(self.date_columns_outer_frame, orient="vertical", command=self.date_columns_canvas.yview)
        self.date_columns_scrollbar.grid(row=0, column=1, sticky='ns')

        self.date_columns_canvas.configure(yscrollcommand=self.date_columns_scrollbar.set)
        self.date_columns_canvas.bind('<Configure>', lambda e: self.date_columns_canvas.configure(scrollregion = self.date_columns_canvas.bbox("all")))
        self.date_columns_canvas.bind_all("<MouseWheel>", self._on_mousewheel_date_format)

        self.date_columns_frame = ttk.Frame(self.date_columns_canvas, padding="5 5 5 5")
        self.date_columns_canvas.create_window((0, 0), window=self.date_columns_frame, anchor="nw")
        self.date_columns_frame.bind("<Configure>", lambda e: self.date_columns_canvas.configure(scrollregion=self.date_columns_canvas.bbox("all")))

        ttk.Label(self.date_columns_frame, text="Load a file to see columns...", background=self.bg_color).pack(pady=10)


        # --- Tab 5: Sorting ---
        sorting_tab = ttk.Frame(self.notebook, padding="15 15 15 15")
        self.notebook.add(sorting_tab, text="5. Sorting")

        sorting_tab.grid_columnconfigure(1, weight=1)
        sorting_tab.grid_rowconfigure(0, weight=1)

        ttk.Label(sorting_tab, text="Sort Data By Column:").grid(row=0, column=0, sticky='w', pady=5)
        self.sort_column_combobox = ttk.Combobox(sorting_tab, textvariable=self.sort_column,
                                                 values=[], state='readonly')
        self.sort_column_combobox.grid(row=0, column=1, padx=5, pady=5, sticky='ew')

        ttk.Label(sorting_tab, text="Sort Order:").grid(row=1, column=0, sticky='w', pady=5)
        sort_order_frame = ttk.Frame(sorting_tab)
        sort_order_frame.grid(row=1, column=1, padx=5, pady=5, sticky='ew')
        ttk.Radiobutton(sort_order_frame, text="Ascending", variable=self.sort_order, value=True).pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(sort_order_frame, text="Descending", variable=self.sort_order, value=False).pack(side=tk.LEFT, padx=5)


        # --- Tab 6: Missing Values & Review ---
        missing_values_tab = ttk.Frame(self.notebook, padding="15 15 15 15")
        self.notebook.add(missing_values_tab, text="6. Missing Values & Review")
        missing_values_tab.grid_columnconfigure(0, weight=1)
        missing_values_tab.grid_rowconfigure(1, weight=1)

        ttk.Label(missing_values_tab, text="Rows with Empty Cells (click cell to edit):",
                  font=('Helvetica', 10, 'bold'), foreground=self.fg_color).grid(row=0, column=0, sticky='w', pady=(0,5))

        # Treeview for missing values
        tree_frame = ttk.Frame(missing_values_tab)
        tree_frame.grid(row=1, column=0, sticky='nsew', pady=5)
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)

        self.missing_values_tree = ttk.Treeview(tree_frame, show='headings')
        self.missing_values_tree.grid(row=0, column=0, sticky='nsew')

        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.missing_values_tree.yview)
        vsb.grid(row=0, column=1, sticky='ns')
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.missing_values_tree.xview)
        hsb.grid(row=1, column=0, sticky='ew')

        self.missing_values_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self.missing_values_tree.bind("<Button-1>", self._on_missing_treeview_click) # Bind click for editing


        # --- Process Button (outside tabs) ---
        self.process_button = ttk.Button(main_frame, text="Process File", command=self.process_file)
        self.process_button.pack(pady=20)

        # --- Progress and Status (outside tabs) ---
        self.progress_status_label = ttk.Label(main_frame, text="Ready.", font=('Helvetica', 10, 'bold'), foreground=self.fg_color)
        self.progress_status_label.pack(pady=(0, 5), fill='x')

        self.progress_bar = ttk.Progressbar(main_frame, orient='horizontal', length=400, mode='determinate')
        self.progress_bar.pack(pady=(0, 10), fill='x')

        self.status_label = ttk.Label(main_frame, text="Select a file and folder to begin.",
                                      font=('Helvetica', 10, 'italic'), wraplength=600)
        self.status_label.pack(pady=(5, 0), fill='x')

        # Try to load initial columns if a file path is already set from config
        if self.input_file_path.get() and os.path.exists(self.input_file_path.get()):
            self._load_dataframe_and_update_widgets(self.input_file_path.get())


    def _on_mousewheel_date_format(self, event):
        """Allows mousewheel scrolling for the date format canvas."""
        self.date_columns_canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    def _on_mousewheel_duplicate_check(self, event):
        """Allows mousewheel scrolling for the duplicate check canvas."""
        self.duplicate_check_columns_canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    def _on_missing_treeview_click(self, event):
        """Handles clicks on the Missing Values Treeview to enable editing."""
        if self.df is None: return

        item = self.missing_values_tree.identify_row(event.y)
        column_id = self.missing_values_tree.identify_column(event.x)

        if not item or not column_id:
            return

        # Get column name from Treeview's internal column ID
        column_index = int(column_id.replace('#', '')) - 1 # Column ID is #1, #2 etc.
        if column_index < 0 or column_index >= len(self.df.columns):
            return # Should not happen if Treeview is correctly populated

        col_name = self.df.columns[column_index]
        
        # The 'item' variable already holds the iid, which is the original DataFrame index (as a string)
        original_df_index = item 
        
        # Check if the actual value in the DataFrame is NaN
        # We only want to open the editor if the cell is currently empty (NaN)
        current_df_value = self.df.loc[int(original_df_index), col_name]

        if pd.isna(current_df_value):
            # Get bounding box of the cell
            x, y, width, height = self.missing_values_tree.bbox(item, column_id)

            # Create an Entry widget over the cell
            entry = ttk.Entry(self.missing_values_tree, style='TEntry') # Use 'TEntry' style
            entry.place(x=x, y=y, width=width, height=height, anchor='nw')
            
            # Insert an empty string if the value is missing (NaN, pd.NA, None, NaT), otherwise convert to string
            entry.insert(0, "" if pd.isna(current_df_value) else str(current_df_value))
            entry.focus_set()

            def save_edit(event=None):
                new_value_str = entry.get() # Get the value as a string from the Entry widget

                # Determine the appropriate value to store in the DataFrame
                value_to_store = pd.NA # Default to missing value (pandas' preferred NaN)

                if new_value_str != "": # If the user typed something
                    try:
                        # Try to convert to integer first
                        value_to_store = int(new_value_str)
                    except ValueError:
                        try:
                            # If integer conversion fails, try float
                            value_to_store = float(new_value_str)
                        except ValueError:
                            # If neither int nor float, keep it as a string
                            value_to_store = new_value_str

                # Update the DataFrame using .at for efficient single-cell assignment
                self.df.at[int(original_df_index), col_name] = value_to_store
                
                # Update the Treeview to reflect the change
                # Retrieve all values for the row from the updated DataFrame
                updated_row_values = self.df.loc[int(original_df_index)].values.tolist()
                
                # Convert NaNs (including pd.NA) in the list to empty strings for Treeview display
                display_values = ["" if pd.isna(val) else val for val in updated_row_values]
                
                self.missing_values_tree.item(item, values=display_values)
                
                entry.destroy() # Remove the Entry widget

            entry.bind("<Return>", save_edit) # Save on Enter key
            entry.bind("<FocusOut>", save_edit) # Save when focus is lost (e.g., clicking away)

    def _on_duplicate_treeview_click(self, event):
        """Handles clicks on the Duplicate Review Treeview to toggle 'Keep' checkbox."""
        if self.duplicate_review_df is None: return

        item = self.duplicate_review_tree.identify_row(event.y)
        column_id = self.duplicate_review_tree.identify_column(event.x)

        if not item or not column_id:
            return

        # Check if the click was on the "Keep" column (which is the first column, index 0)
        if int(column_id.replace('#', '')) == 1: # Column index in Treeview is 1-based
            original_df_index = int(item) # The iid is the original DF index

            # Toggle the BooleanVar
            current_keep_state = self.duplicate_rows_to_keep_vars[original_df_index].get()
            self.duplicate_rows_to_keep_vars[original_df_index].set(not current_keep_state)

            # Update the Treeview checkbox display
            new_display_value = "✓" if not current_keep_state else "☐"
            values = list(self.duplicate_review_tree.item(item, 'values'))
            values[0] = new_display_value # Assuming "Keep" is the first column
            self.duplicate_review_tree.item(item, values=values)

    def browse_input_file(self):
        initial_dir = os.path.dirname(self.input_file_path.get()) if self.input_file_path.get() else os.path.expanduser('~')
        file_types = [("Excel files", "*.xlsx"), ("CSV files", "*.csv"), ("All files", "*.*")]
        file_path = filedialog.askopenfilename(
            title="Select Input File",
            initialdir=initial_dir,
            filetypes=file_types
        )
        if file_path:
            self.input_file_path.set(file_path)
            self.update_status(f"Input file selected: {os.path.basename(file_path)}")
            self._save_config()
            self._load_dataframe_and_update_widgets(file_path)

    def browse_output_folder(self):
        initial_dir = self.output_folder_path.get() if self.output_folder_path.get() else os.path.expanduser('~')
        folder_path = filedialog.askdirectory(
            title="Select Output Folder",
            initialdir=initial_dir
        )
        if folder_path:
            self.output_folder_path.set(folder_path)
            self.update_status(f"Output folder selected: {os.path.basename(folder_path)}")
            self._save_config()

    def _load_dataframe_and_update_widgets(self, file_path):
        """Loads the entire DataFrame and then updates all column-dependent widgets."""
        self.update_progress_status("Loading entire file...", 0)
        try:
            if file_path.lower().endswith('.csv'):
                self.df = pd.read_csv(file_path)
            elif file_path.lower().endswith('.xlsx'):
                self.df = pd.read_excel(file_path)
            else:
                self.update_progress_status("Unsupported file type for loading.", 0)
                messagebox.showerror("File Type Error", "Unsupported file type. Please select a .csv or .xlsx file.")
                self.df = None
                return

            if self.df.empty:
                messagebox.showwarning("Empty File", "The selected file is empty.")
                self.update_progress_status("Empty file loaded.", 0)
                self.df = None
                return

            columns = self.df.columns.tolist()
            
            # --- Update Date Column Checkboxes ---
            for widget in self.date_columns_frame.winfo_children():
                widget.destroy()
            self.date_column_vars.clear()
            
            if not columns:
                 ttk.Label(self.date_columns_frame, text="No columns found in file.", background=self.bg_color).pack(pady=10)
            else:
                col_idx = 0
                for col in columns:
                    var = tk.BooleanVar()
                    if any(keyword in col.lower() for keyword in ["date", "day", "month"]):
                        var.set(True)
                    self.date_column_vars[col] = var
                    chk = ttk.Checkbutton(self.date_columns_frame, text=col, variable=var)
                    chk.grid(row=col_idx // 3, column=col_idx % 3, sticky='w', padx=2, pady=2)
                    col_idx += 1
                self.date_columns_canvas.update_idletasks()
                self.date_columns_canvas.config(scrollregion=self.date_columns_canvas.bbox("all"))

            # --- Update Duplicate Check Column Checkboxes ---
            for widget in self.duplicate_check_columns_frame.winfo_children():
                widget.destroy()
            self.duplicate_check_column_vars.clear()

            if not columns:
                 ttk.Label(self.duplicate_check_columns_frame, text="No columns found in file.", background=self.bg_color).pack(pady=10)
            else:
                col_idx = 0
                for col in columns:
                    var = tk.BooleanVar(value=True) # Default to all columns selected for duplicate check
                    self.duplicate_check_column_vars[col] = var
                    chk = ttk.Checkbutton(self.duplicate_check_columns_frame, text=col, variable=var)
                    chk.grid(row=col_idx // 3, column=col_idx % 3, sticky='w', padx=2, pady=2)
                    col_idx += 1
                self.duplicate_check_columns_canvas.update_idletasks()
                self.duplicate_check_columns_canvas.config(scrollregion=self.duplicate_check_columns_canvas.bbox("all"))

            # --- Update Missing Values Treeview ---
            self._populate_missing_values_treeview()

            # --- Clear Duplicate Review Treeview (reset for new file) ---
            self._clear_duplicate_review_treeview()
            self.apply_duplicates_button.config(state='disabled')

            # --- Update Sort Column Combobox ---
            self.sort_column_combobox['values'] = columns
            if columns:
                if self.sort_column.get() not in columns:
                    self.sort_column.set(columns[0])
            else:
                self.sort_column.set("")

            self.update_progress_status("File loaded and options updated.", 10)

        except Exception as e:
            messagebox.showerror("Error Loading File", f"Could not load or read file: {e}")
            self.update_progress_status("Error loading file.", 0)
            self.df = None
            # Clear all dependent widgets
            for widget in self.date_columns_frame.winfo_children(): widget.destroy()
            self.date_column_vars.clear()
            for widget in self.duplicate_check_columns_frame.winfo_children(): widget.destroy()
            self.duplicate_check_column_vars.clear()
            self._clear_missing_values_treeview()
            self._clear_duplicate_review_treeview()
            self.sort_column_combobox['values'] = []
            self.sort_column.set("")

    def _populate_missing_values_treeview(self):
        """Populates the Treeview with rows containing NaN values."""
        self._clear_missing_values_treeview()

        if self.df is None or self.df.empty:
            return

        # Identify rows with any NaN values
        missing_rows_df = self.df[self.df.isnull().any(axis=1)]

        if missing_rows_df.empty:
            self.missing_values_tree.heading("#0", text="")
            self.missing_values_tree["columns"] = ()
            # Placeholder for "No empty cells" message
            ttk.Label(self.missing_values_tree.master, text="No empty cells found in file.", background=self.bg_color, foreground=self.fg_color).pack(pady=10)
            return

        # Define Treeview columns
        columns = self.df.columns.tolist()
        self.missing_values_tree["columns"] = columns
        self.missing_values_tree.column("#0", width=0, stretch=tk.NO)
        
        for col in columns:
            self.missing_values_tree.heading(col, text=col)
            self.missing_values_tree.column(col, width=100, anchor='w')

        # Insert data
        for original_idx, row in missing_rows_df.iterrows():
            display_values = ["" if pd.isna(val) else val for val in row.values.tolist()]
            self.missing_values_tree.insert("", "end", iid=str(original_idx), values=display_values)

    def _clear_missing_values_treeview(self):
        """Clears all data from the missing values Treeview."""
        if self.missing_values_tree:
            for item in self.missing_values_tree.get_children():
                self.missing_values_tree.delete(item)
            self.missing_values_tree["columns"] = ()
            self.missing_values_tree.heading("#0", text="")
            # Remove any "No empty cells" label if present
            for widget in self.missing_values_tree.master.winfo_children():
                if isinstance(widget, ttk.Label) and widget.cget("text") == "No empty cells found in file.":
                    widget.destroy()

    def _populate_duplicate_review_treeview(self):
        """Populates the Treeview with duplicate rows for review."""
        self._clear_duplicate_review_treeview()

        if self.duplicate_review_df is None or self.duplicate_review_df.empty:
            self.duplicate_review_tree.heading("#0", text="")
            self.duplicate_review_tree["columns"] = ()
            ttk.Label(self.duplicate_review_tree.master, text="No duplicates found for review.", background=self.bg_color, foreground=self.fg_color).pack(pady=10)
            self.apply_duplicates_button.config(state='disabled')
            return

        # Define Treeview columns: "Keep" checkbox + original DataFrame columns
        columns_to_display = ["Keep"] + self.duplicate_review_df.columns.tolist()
        self.duplicate_review_tree["columns"] = columns_to_display
        self.duplicate_review_tree.column("#0", width=0, stretch=tk.NO) # Hide default first column
        
        for col in columns_to_display:
            self.duplicate_review_tree.heading(col, text=col)
            if col == "Keep":
                self.duplicate_review_tree.column(col, width=50, anchor='center')
            else:
                self.duplicate_review_tree.column(col, width=100, anchor='w')

        self.duplicate_rows_to_keep_vars.clear()
        
        # Group duplicates to ensure at least one is kept from each group
        # This uses the same subset_cols that were used to find the duplicates
        subset_cols = [col_name for col_name, var in self.duplicate_check_column_vars.items() if var.get()]
        
        # Create a temporary DataFrame to find groups of duplicates
        # Use the original (pre-cleaned) df for grouping to ensure correct original indices are used
        temp_df = self.df.copy() 
        
        # Identify group IDs for all rows that are part of a duplicate set
        # This will assign a unique ID to each group of identical rows (based on subset_cols)
        # Rows that are not duplicates will get NaN, which is fine as we filter later.
        temp_df['__dup_group_id'] = temp_df.groupby(subset_cols, dropna=False).ngroup()

        # Iterate through unique group IDs that have duplicates
        # We only care about groups where more than one row exists
        duplicate_group_ids = temp_df[temp_df.duplicated(subset=subset_cols, keep=False)]['__dup_group_id'].unique()

        for group_id in duplicate_group_ids:
            group_rows = temp_df[temp_df['__dup_group_id'] == group_id]
            
            # For each group, we'll keep the first one by default, and mark others for deletion
            first_in_group = True
            for original_idx, row_data in group_rows.iterrows():
                var = tk.BooleanVar(value=True if first_in_group else False) # Keep first by default
                self.duplicate_rows_to_keep_vars[original_idx] = var
                
                display_keep = "✓" if var.get() else "☐"
                # Ensure we display the original values from the self.df.columns, not temp_df's extra column
                display_values = [display_keep] + ["" if pd.isna(val) else val for val in row_data[self.df.columns].values.tolist()]
                
                self.duplicate_review_tree.insert("", "end", iid=str(original_idx), values=display_values)
                first_in_group = False

        self.apply_duplicates_button.config(state='normal')

    def _clear_duplicate_review_treeview(self):
        """Clears all data from the duplicate review Treeview."""
        if self.duplicate_review_tree:
            for item in self.duplicate_review_tree.get_children():
                self.duplicate_review_tree.delete(item)
            self.duplicate_review_tree["columns"] = ()
            self.duplicate_review_tree.heading("#0", text="")
            for widget in self.duplicate_review_tree.master.winfo_children():
                if isinstance(widget, ttk.Label) and widget.cget("text") == "No duplicates found for review.":
                    widget.destroy()
            self.duplicate_rows_to_keep_vars.clear()

    def _apply_duplicate_changes(self):
        """Applies user's duplicate review choices and continues processing."""
        self.update_progress_status("Applying duplicate changes...", 0)
        self.root.update_idletasks()

        rows_to_keep_indices = [idx for idx, var in self.duplicate_rows_to_keep_vars.items() if var.get()]
        
        # --- Capture deleted indices for the report ---
        deleted_indices = []
        for original_idx, var in self.duplicate_rows_to_keep_vars.items():
            if not var.get(): # If the user chose NOT to keep this row
                deleted_indices.append(original_idx)
        self.processing_metrics['deleted_duplicate_original_indices'] = sorted(deleted_indices)

        # --- Create a temporary DataFrame with group IDs for validation ---
        subset_cols = [col_name for col_name, var in self.duplicate_check_column_vars.items() if var.get()]
        
        # Use a copy of the *original* self.df to ensure correct grouping logic
        # dropna=False ensures that rows with NaNs in subset_cols are also grouped
        temp_df_with_groups = self.df.copy() 
        temp_df_with_groups['__dup_group_id'] = temp_df_with_groups.groupby(subset_cols, dropna=False).ngroup()

        # Identify all group IDs that were originally part of a duplicate set
        original_duplicate_group_ids = temp_df_with_groups[temp_df_with_groups.duplicated(subset=subset_cols, keep=False)]['__dup_group_id'].unique()

        # Get the group IDs of the rows that the user chose to keep
        # We filter temp_df_with_groups by the indices the user chose to keep
        kept_rows_with_groups = temp_df_with_groups.loc[rows_to_keep_indices]
        kept_group_ids = kept_rows_with_groups['__dup_group_id'].unique() if not kept_rows_with_groups.empty else []

        all_groups_covered = True
        for group_id in original_duplicate_group_ids:
            if group_id not in kept_group_ids:
                all_groups_covered = False
                break
        
        if not all_groups_covered:
            messagebox.showwarning("Duplicate Review Error", "At least one row from each duplicate group must be kept. Please review again.")
            self.update_progress_status("Error: Review incomplete.", 0)
            # Clear the deleted indices if validation fails, as the changes aren't applied yet
            self.processing_metrics['deleted_duplicate_original_indices'] = [] 
            return

        # --- Construct the final DataFrame after review ---
        # Get rows that were never duplicates (based on the chosen subset_cols)
        non_duplicate_rows = self.df[~self.df.duplicated(subset=subset_cols, keep=False)]
        
        # Get the rows that the user explicitly chose to keep from duplicate sets
        rows_kept_from_duplicates = self.df.loc[rows_to_keep_indices]

        # Combine non-duplicate rows with the ones chosen to be kept from duplicate sets
        # Use pd.concat and then drop duplicates based on the index to handle potential overlaps
        final_processed_df = pd.concat([non_duplicate_rows, rows_kept_from_duplicates]).drop_duplicates(keep='first').reset_index(drop=True)
        
        # Store metrics
        # The number of duplicates removed is the difference between original rows and final rows
        self.processing_metrics['num_duplicates_removed'] = self.processing_metrics['original_rows'] - len(final_processed_df)
        self.processing_metrics['duplicates_reviewed'] = True
        
        self.update_progress_status("Duplicate review applied. Continuing processing...", 0)
        self.root.update_idletasks()
        
        self._final_processing_steps(final_processed_df)


    def process_file(self):
        # Use self.df directly which has been loaded and potentially edited
        if self.df is None:
            messagebox.showwarning("No File Loaded", "Please load an input file first.")
            self.update_progress_status("Error: No file loaded.", 0)
            return

        # Reset metrics for a new run
        self.processing_metrics = {
            'original_rows': len(self.df),
            'whitespace_trimmed': self.do_trim_whitespace.get(),
            'capitalization_applied': self.do_capitalize_strings.get(),
            'date_formatted_cols': [],
            'duplicate_removal_enabled': self.do_remove_duplicates.get(),
            'duplicate_check_cols': [],
            'num_duplicates_removed': 0, # Will be updated later
            'duplicates_reviewed': False, # Will be updated later
            'deleted_duplicate_original_indices': [], # NEW: To store original indices of deleted duplicates
            'sorted_by': None,
            'sorted_order': None,
            'final_rows': 0
        }

        # --- Validation ---
        output_folder = self.output_folder_path.get()
        selected_strftime_format = self._date_format_map.get(self.selected_date_format_display.get())
        
        if not output_folder:
            messagebox.showwarning("Output Error", "Please select an output folder.")
            self.update_progress_status("Error: No output folder selected.", 0)
            return
        if not selected_strftime_format:
            messagebox.showwarning("Date Format Error", "Please select a valid date format from the dropdown.")
            self.update_progress_status("Error: No date format selected.", 0)
            return
        if self.do_remove_duplicates.get():
            selected_duplicate_columns = [col_name for col_name, var in self.duplicate_check_column_vars.items() if var.get()]
            if not selected_duplicate_columns:
                messagebox.showwarning("Duplicate Check Error", "If 'Remove Duplicate Rows' is checked, you must select at least one column for duplicate checking.")
                self.update_progress_status("Error: No columns selected for duplicate check.", 0)
                return
            self.processing_metrics['duplicate_check_cols'] = selected_duplicate_columns

        self.update_progress_status("Starting processing...", 0)
        self.root.update_idletasks()

        try:
            # Create a copy of the DataFrame to apply initial cleaning steps
            processed_df_stage1 = self.df.copy()
            current_progress = 20

            # --- Step 1: Trim Whitespace (Conditional) ---
            if self.do_trim_whitespace.get():
                self.update_progress_status("Trimming whitespace...", current_progress)
                for col in processed_df_stage1.select_dtypes(include=['object']).columns:
                    processed_df_stage1[col] = processed_df_stage1[col].str.strip()
            current_progress += 10

            # --- Step 2: Capitalize String Columns (Conditional) ---
            if self.do_capitalize_strings.get():
                self.update_progress_status("Capitalizing text fields...", current_progress)
                for col in processed_df_stage1.select_dtypes(include=['object']).columns:
                    processed_df_stage1[col] = processed_df_stage1[col].str.title()
            current_progress += 10

            # --- Step 3: Duplicate Review (Interactive) ---
            if self.do_remove_duplicates.get():
                self.update_progress_status("Checking for duplicates...", current_progress)
                subset_cols = [col_name for col_name, var in self.duplicate_check_column_vars.items() if var.get()]
                
                # Find all rows involved in a duplicate set
                self.duplicate_review_df = processed_df_stage1[processed_df_stage1.duplicated(subset=subset_cols, keep=False)].copy()

                if not self.duplicate_review_df.empty:
                    self.update_progress_status(f"{len(self.duplicate_review_df)} rows involved in duplicates. Review required.", current_progress)
                    self._populate_duplicate_review_treeview()
                    self.notebook.select(2) # Switch to Duplicate Review tab
                    self.process_button.config(state='disabled') # Disable main process button
                    self.apply_duplicates_button.config(state='normal') # Enable apply button
                    return # Pause execution until user reviews
                else:
                    self.update_status("No duplicates found for review. Proceeding...")
                    self.processing_metrics['num_duplicates_removed'] = 0
                    self.processing_metrics['duplicates_reviewed'] = True # Mark as reviewed (even if empty)
                    self.processing_metrics['deleted_duplicate_original_indices'] = [] # No duplicates, so none deleted
            
            # If no duplicates to review or duplicate removal is off, proceed directly
            self.update_status("Proceeding to final processing steps...")
            self._final_processing_steps(processed_df_stage1)

        except Exception as e:
            messagebox.showerror("An Error Occurred During Initial Processing", f"An unexpected error occurred: {e}")
            self.update_progress_status(f"Error: {e}", 0)
            self.process_button.config(state='normal') # Re-enable button on error
            self.apply_duplicates_button.config(state='disabled') # Disable apply button on error
        finally:
            self.progress_bar['value'] = 0
            self.root.update_idletasks()

    def _final_processing_steps(self, processed_df):
        """Applies date formatting, sorting, saves file, and generates report."""
        try:
            current_progress = 60 # Starting after initial cleaning and duplicate handling
            original_rows_for_report = self.processing_metrics['original_rows'] # Use original from metrics

            # --- Step 4: Date Column Processing ---
            self.update_progress_status("Formatting selected date columns...", current_progress)
            processed_date_cols = []
            selected_strftime_format = self._date_format_map.get(self.selected_date_format_display.get())
            for col_name, var in self.date_column_vars.items():
                if var.get() and col_name in processed_df.columns:
                    processed_date_cols.append(col_name)
                    try:
                        processed_df[col_name] = pd.to_datetime(processed_df[col_name], errors='coerce')
                        processed_df[col_name] = processed_df[col_name].dt.strftime(selected_strftime_format).fillna('')
                    except Exception as e:
                        messagebox.showwarning("Date Formatting Warning",
                                               f"Could not apply format to column '{col_name}'. "
                                               f"Error: {e}\nInvalid dates will be cleared.")
                        processed_df[col_name] = processed_df[col_name].fillna('')
            self.processing_metrics['date_formatted_cols'] = processed_date_cols
            current_progress += 10

            # --- Step 5: Sort Data ---
            sort_col = self.sort_column.get()
            sort_asc = self.sort_order.get()
            if sort_col and sort_col in processed_df.columns:
                self.update_progress_status(f"Sorting data by '{sort_col}'...", current_progress)
                processed_df = processed_df.sort_values(by=sort_col, ascending=sort_asc)
                self.processing_metrics['sorted_by'] = sort_col
                self.processing_metrics['sorted_order'] = 'Ascending' if sort_asc else 'Descending'
            current_progress += 10

            self.processing_metrics['final_rows'] = len(processed_df)

            # --- Step 6: Save Cleaned DataFrame ---
            self.update_progress_status("Saving cleaned file...", current_progress)
            original_filename_with_ext = os.path.basename(self.input_file_path.get())
            original_name_without_ext, _ = os.path.splitext(original_filename_with_ext)
            
            output_ext = self.output_file_format.get()
            cleaned_filename = f"cleaned.{original_name_without_ext}{output_ext}"
            output_file_path = os.path.join(self.output_folder_path.get(), cleaned_filename)

            if output_ext == ".xlsx":
                processed_df.to_excel(output_file_path, index=False)
            elif output_ext == ".csv":
                processed_df.to_csv(output_file_path, index=False)
            current_progress += 10

            # --- Step 7: Generate Processing Report (Conditional) ---
            if self.generate_report.get():
                self.update_progress_status("Generating processing report...", current_progress)
                self._generate_processing_report(output_file_path, original_name_without_ext)
            current_progress += 10
            
            self.update_progress_status("Processing complete!", 100)

            # --- Final Summary Message ---
            summary_message = (
                f"Successfully processed:\n"
                f"Original rows: {original_rows_for_report}\n"
            )
            if self.processing_metrics['duplicates_reviewed']: # Was duplicate removal enabled and/or reviewed?
                 if self.processing_metrics['num_duplicates_removed'] > 0:
                     summary_message += f"Duplicates removed: {self.processing_metrics['num_duplicates_removed']}\n"
                 else:
                     summary_message += "No duplicates removed (or found).\n"
            else: # Duplicate removal was skipped entirely
                summary_message += "Duplicate removal was skipped.\n"
            
            summary_message += (
                f"Final rows: {self.processing_metrics['final_rows']}\n"
                f"Cleaned file saved to:\n{output_file_path}"
            )
            if self.generate_report.get():
                report_name = f"{original_name_without_ext}_report.docx"
                summary_message += f"\nProcessing report saved as:\n{os.path.join(self.output_folder_path.get(), report_name)}"

            messagebox.showinfo("Processing Complete", summary_message)

        except Exception as e:
            messagebox.showerror("An Error Occurred During Final Processing", f"An unexpected error occurred: {e}")
            self.update_progress_status(f"Error: {e}", 0)
        finally:
            self.process_button.config(state='normal') # Re-enable main process button
            self.apply_duplicates_button.config(state='disabled') # Ensure apply button is disabled
            self.progress_bar['value'] = 0
            self.root.update_idletasks()

    def _generate_processing_report(self, cleaned_file_path, original_name_without_ext):
        """Generates a Word document report of the cleaning process."""
        try:
            document = Document()
            
            # Title
            document.add_heading(f"Processing Report for '{original_name_without_ext}'", level=1)
            document.add_paragraph(f"Generated by CleanIt on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            document.add_paragraph("---")

            # Summary Section
            document.add_heading("Summary", level=2)
            document.add_paragraph(f"Original Rows: {self.processing_metrics.get('original_rows', 'N/A')}")
            document.add_paragraph(f"Final Rows: {self.processing_metrics.get('final_rows', 'N/A')}")
            document.add_paragraph(f"Cleaned File Saved To: {cleaned_file_path}")
            document.add_paragraph("---")

            # Cleaning Operations Details
            document.add_heading("Cleaning Operations Performed", level=2)

            document.add_paragraph(f"Whitespace Trimming: {'Enabled' if self.processing_metrics.get('whitespace_trimmed') else 'Disabled'}")
            document.add_paragraph(f"Text Field Capitalization: {'Enabled' if self.processing_metrics.get('capitalization_applied') else 'Disabled'}")
            
            dup_status = "Enabled" if self.processing_metrics.get('duplicate_removal_enabled') else "Disabled"
            document.add_paragraph(f"Duplicate Removal: {dup_status}")
            if self.processing_metrics.get('duplicate_removal_enabled'):
                document.add_paragraph(f"  - Columns used for duplicate check: {', '.join(self.processing_metrics.get('duplicate_check_cols', []))}")
                document.add_paragraph(f"  - Duplicates Removed: {self.processing_metrics.get('num_duplicates_removed', 0)}")
                document.add_paragraph(f"  - Interactive Review Performed: {'Yes' if self.processing_metrics.get('duplicates_reviewed') else 'No (skipped due to no duplicates)'}")
                
                # NEW: Add deleted duplicate row indices
                deleted_dup_indices = self.processing_metrics.get('deleted_duplicate_original_indices', [])
                if deleted_dup_indices:
                    # Convert list of integers to string for display, e.g., "1, 5, 10-12"
                    # For simplicity, just list them, but could make it fancier
                    indices_str = ", ".join(map(str, deleted_dup_indices))
                    document.add_paragraph(f"  - Original row indices of deleted duplicates: {indices_str}")
                else:
                    document.add_paragraph("  - No specific rows deleted as duplicates (either none found or all kept during review).")


            date_cols = self.processing_metrics.get('date_formatted_cols', [])
            date_format_used = self.selected_date_format_display.get()
            document.add_paragraph(f"Date Formatting: {'Applied' if date_cols else 'Not Applied'}")
            if date_cols:
                document.add_paragraph(f"  - Columns Formatted: {', '.join(date_cols)}")
                document.add_paragraph(f"  - Format Applied: '{date_format_used}'")

            sort_col = self.processing_metrics.get('sorted_by')
            if sort_col:
                document.add_paragraph(f"Data Sorted By: '{sort_col}' ({self.processing_metrics.get('sorted_order')})")
            else:
                document.add_paragraph("Data Sorting: Not Applied")
            
            # Save the document
            report_filename = f"{original_name_without_ext}_report.docx"
            report_path = os.path.join(self.output_folder_path.get(), report_filename)
            document.save(report_path)
            self.update_status(f"Processing report saved to: {report_path}")

        except Exception as e:
            messagebox.showerror("Report Generation Error", f"Could not generate processing report: {e}")
            self.update_status(f"Error generating report: {e}")


    def update_progress_status(self, message, value):
        self.progress_status_label.config(text=message)
        self.progress_bar['value'] = value
        self.root.update_idletasks()

    def update_status(self, message):
        self.status_label.config(text=message)

# --- Main execution ---
if __name__ == "__main__":
    root = tk.Tk()
    app = CleanItApp(root)
    root.mainloop()
