#make sure you use p_o_p and not p.o.p
#use sitrep_data: common error


#auto install packages for end users
import subprocess
import sys

def install_packages():
    """Installs required packages using pip."""
    packages = ['python-docx', 'ttkthemes']

    try:
        for package in packages:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    except subprocess.CalledProcessError as e:
        print(f"Error installing packages: {e}")
        print("Please make sure you have pip installed and try again.")
        sys.exit(1)  # Exit the script if installation fails


# Call the install_packages function before anything else
#install_packages()

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import re  
import calendar
from tkcalendar import Calendar

class SitrepGenerator:
    def __init__(self, master):
        self.master = master
        master.title("SITREP Generator")

        # Color Scheme
        self.bg_color = "#E6F7FF"  # Light Blue (background)
        self.fg_color = "#000080"  # Dark Blue (for text)
        master.configure(bg=self.bg_color)

        # Style for ttk widgets
        style = ttk.Style()
        style.configure("TNotebook", background=self.bg_color)
        style.configure("TFrame", background=self.bg_color)
        style.configure("TLabel", background=self.bg_color, foreground=self.fg_color)
        style.configure("TButton", background="#B0E0E6", foreground=self.fg_color)  # Powder Blue

        # Fullscreen
        master.state('zoomed')  # or master.attributes('-fullscreen', True)

        # Notebook (Tabs)
        self.notebook = ttk.Notebook(master)
        self.sitrep_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.sitrep_tab, text="SITREP")
        self.notebook.pack(expand=True, fill="both")

        # Variables
        self.output_folder_path = tk.StringVar()
        self.output_filename = tk.StringVar()
        self.sitrep_data = []  # List to store sitrep data
        self.prompt_type = tk.StringVar(value="Award")  # Default prompt type
        self.custom_prompt = tk.StringVar()  # Variable for custom prompt text
        self.selected_sitrep_index = None  # Track the index of the selected SITREP

        # Product List (Alphabetical Order)
        self.product_list = sorted([
            "Air",
            "Argon",
            "Carbon Dioxide",
            "Deuterium",
            "Dinitrogen Tetroxide",
            "Fluorine",
            "Helium",
            "Hydrazine Aerozine",
            "Hydrazine Anhydrous",
            "Hydrazine High Purity",
            "Hydrazine Monomethyl (MMH)",
            "Hydrazine Monopropellant Anhydrous",
            "Hydrazine Unsymmetrical-Dimethyl",
            "Hydrazine-Water (H-70)",
            "Hydrogen",
            "Hydrogen Peroxide",
            "Isopropyl Alcohol",
            "JP-10 High Density Synthetic Hydrocarbon",
            "JP-7 Turbine Fuel",
            "Krypton",
            "MAF-4",
            "Methane",
            "Methanol",
            "Neon",
            "Nitric Acid",
            "Nitrogen",
            "Oxygen",
            "Priming Fluid (PF-1)",
            "RP-1 Kerosene Bulk",
            "RP-1 Kerosene Drum",
            "RP-2 Kerosene Bulk",
            "RP-2 Kerosene Drum",
            "Xenon"
        ])


        # Sitrep Title Label
        tk.Label(self.sitrep_tab, text="SITREP", font=("Arial", 16, "bold"), bg=self.bg_color, fg=self.fg_color).grid(row=0, column=0, columnspan=3, sticky=tk.N, padx=5, pady=10)

        # Labels and Entry Fields (SITREP Tab)
        tk.Label(self.sitrep_tab, text="Output Folder:", bg=self.bg_color, fg=self.fg_color).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.output_entry = tk.Entry(self.sitrep_tab, textvariable=self.output_folder_path, width=60)
        self.output_entry.grid(row=1, column=1, padx=5, pady=5)
        tk.Button(self.sitrep_tab, text="Browse Output", command=self.browse_output).grid(row=1, column=2, padx=5, pady=5)

        tk.Label(self.sitrep_tab, text="Output Filename:", bg=self.bg_color, fg=self.fg_color).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        self.filename_entry = tk.Entry(self.sitrep_tab, textvariable=self.output_filename, width=60)
        self.filename_entry.grid(row=2, column=1, padx=5, pady=5)

        # Prompt Type Dropdown
        tk.Label(self.sitrep_tab, text="Prompt Type:", bg=self.bg_color, fg=self.fg_color).grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        self.prompt_dropdown = ttk.Combobox(self.sitrep_tab, textvariable=self.prompt_type, values=["Award", "CPARS", "Custom", "Data Cleansing", "FPDS-NG", "Issues Solicitation / Synopsis/ Sources Sought Notice", "Negotiation", "Purchase"], state="readonly", width=57)
        self.prompt_dropdown.grid(row=3, column=1, padx=5, pady=5)
        self.prompt_dropdown.bind("<<ComboboxSelected>>", self.update_prompt_fields)

        # Custom Prompt Textbox
        self.custom_prompt_label = tk.Label(self.sitrep_tab, text="SITREP:", bg=self.bg_color, fg=self.fg_color)
        self.custom_prompt_entry = tk.Text(self.sitrep_tab, wrap=tk.WORD, width=60, height=3, bg="white", fg="black")

        # Save Sitrep Button
        self.save_sitrep_button = tk.Button(self.sitrep_tab, text="Save SITREP", command=self.save_sitrep)
        # Generate Button
        self.generate_sitrep_button = tk.Button(self.sitrep_tab, text="Generate SITREP", command=self.generate_sitrep)

        # Sitrep Input Frame
        self.sitrep_input_frame = ttk.Frame(self.sitrep_tab)
        self.sitrep_input_frame.grid(row=5, column=0, columnspan=3, sticky=tk.W + tk.E, padx=5, pady=5)
        self.create_sitrep_input_fields(self.sitrep_input_frame)

        # Text box to display filled templates
        self.template_display = tk.Text(self.sitrep_tab, wrap=tk.WORD, width=45, height=10, bg="white", fg="black", state="disabled") # Set state to disabled
        self.template_display.grid(row=1, column=3, rowspan=7, padx=(5, 20), pady=5, sticky=tk.N + tk.S + tk.E + tk.W)
        self.template_display.bind("<Button-1>", self.select_sitrep) # Bind left mouse click
        # Added padx to right to push it further left

        # Scrollbar for the text box
        self.template_scrollbar = ttk.Scrollbar(self.sitrep_tab, orient=tk.VERTICAL, command=self.template_display.yview)
        self.template_scrollbar.grid(row=1, column=4, rowspan=7, sticky=tk.N + tk.S)
        self.template_display['yscrollcommand'] = self.template_scrollbar.set

        # Delete Button for the Text Box
        self.delete_button = tk.Button(self.sitrep_tab, text="Delete", command=self.delete_selected_sitrep)
        self.delete_button.grid(row=8, column=3, padx=5, pady=5, sticky=tk.E)  # Place next to the text box

        self.update_prompt_fields() # Initialize UI based on default prompt type


    def popup_calendar(self, entry):
        """Creates a popup calendar to select a date with custom styling."""
        def set_date():
            """Sets the date in the entry and closes the popup."""
            try:
                from datetime import datetime
                selected_date_str = cal.get_date()
                selected_date = datetime.strptime(selected_date_str, "%m-%d-%Y") # Parse the string
                formatted_date = selected_date.strftime("%B %d, %Y") # Format the datetime object
                entry.delete(0, tk.END)  # Clear the entry first
                entry.insert(0, formatted_date)  # Insert the formatted date
                top.destroy()
            except ValueError:
                messagebox.showerror("Error", "Invalid date format. Please use the calendar.")

        top = tk.Toplevel(self.master)
        # Customize the calendar colors here
        cal = Calendar(top, 
                       selectmode="day", 
                       date_pattern="mm-dd-yyyy",
                       background="midnightblue",  # Background color
                       foreground="white",    # Date text color
                       normalbackground="white", #background color of normal days
                       normalforeground="black",
                       selectbackground="lightblue", #background color of selected day
                       selectforeground="black",
                       weekendbackground="white",
                       weekendforeground="black",
                       headersbackground="midnightblue",
                       headersforeground="white"
                       )
        cal.pack(pady=10)
        tk.Button(top, text="Set Date", command=set_date).pack()


    def browse_output(self):
        folder = filedialog.askdirectory()
        self.output_folder_path.set(folder)

    def create_sitrep_input_fields(self, frame):
        # Destroy existing widgets in the frame
        for widget in frame.winfo_children():
            widget.destroy()

        if self.prompt_type.get() == "Purchase":
            # Sample Prompt and Input Fields
            prompt_text = "(U) Aerospace Energy [Does What] for [Product] to Support the [Place]: On [Date], Aerospace Energy [Did What] for solicitation [Specific Product] to support [Place] customers with [Customer Support Detail]. The procurement will result in [Term Length], firm fixed-price requirements type contracts with a period of performance from [Time Frame]."
            tk.Label(frame, text="Prompt:", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
            tk.Label(frame, text=prompt_text, wraplength=700, justify="left", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=1, columnspan=2, sticky=tk.W, padx=5, pady=5)

            self.does_what_var = tk.StringVar()
            tk.Label(frame, text="Does What:", bg=self.bg_color, fg=self.fg_color).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
            self.does_what_entry = tk.Entry(frame, textvariable=self.does_what_var, width=50)
            self.does_what_entry.grid(row=1, column=1, padx=5, pady=5)
            self.does_what_var.trace_add('write', self.update_template_display)  # Trace variable changes

            # Replace Item Entry with Dropdown
            self.item_var = tk.StringVar()
            tk.Label(frame, text="Product:", bg=self.bg_color, fg=self.fg_color).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
            self.item_dropdown = ttk.Combobox(frame, textvariable=self.item_var, values=self.product_list, state="readonly", width=47) # Adjusted width
            self.item_dropdown.grid(row=2, column=1, padx=5, pady=5)
            self.item_var.trace_add('write', self.update_template_display)  # Trace variable changes


            self.place_var = tk.StringVar()
            tk.Label(frame, text="Place:", bg=self.bg_color, fg=self.fg_color).grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
            self.place_entry = tk.Entry(frame, textvariable=self.place_var, width=50)
            self.place_entry.grid(row=3, column=1, padx=5, pady=5)
            self.place_var.trace_add('write', self.update_template_display)  # Trace variable changes

            self.date_value_var = tk.StringVar()
            tk.Label(frame, text="Date:", bg=self.bg_color, fg=self.fg_color).grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
            self.date_value_entry = tk.Entry(frame, textvariable=self.date_value_var, width=50)
            self.date_value_entry.grid(row=4, column=1, padx=5, pady=5)
            self.date_value_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.date_value_entry))
            self.date_value_var.trace_add('write', self.update_template_display)  # Trace variable changes

            self.did_what_var = tk.StringVar()
            tk.Label(frame, text="Did What:", bg=self.bg_color, fg=self.fg_color).grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
            self.did_what_entry = tk.Entry(frame, textvariable=self.did_what_var, width=50)
            self.did_what_entry.grid(row=5, column=1, padx=5, pady=5)
            self.did_what_var.trace_add('write', self.update_template_display)  # Trace variable changes

            self.specific_item_var = tk.StringVar()
            tk.Label(frame, text="Specific Product:", bg=self.bg_color, fg=self.fg_color).grid(row=6, column=0, sticky=tk.W, padx=5, pady=5)
            self.specific_item_entry = tk.Entry(frame, textvariable=self.specific_item_var, width=50)
            self.specific_item_entry.grid(row=6, column=1, padx=5, pady=5)
            self.specific_item_var.trace_add('write', self.update_template_display)  # Trace variable changes

            self.customer_support_detail_var = tk.StringVar()
            tk.Label(frame, text="Customer Support Detail:", bg=self.bg_color, fg=self.fg_color).grid(row=7, column=0, sticky=tk.W, padx=5, pady=5)
            self.customer_support_detail_entry = tk.Entry(frame, textvariable=self.customer_support_detail_var, width=50)
            self.customer_support_detail_entry.grid(row=7, column=1, padx=5, pady=5)
            self.customer_support_detail_var.trace_add('write', self.update_template_display)  # Trace variable changes

            self.time_frame_var = tk.StringVar()
            tk.Label(frame, text="Time Frame:", bg=self.bg_color, fg=self.fg_color).grid(row=8, column=0, sticky=tk.W, padx=5, pady=5)
            self.time_frame_entry = tk.Entry(frame, textvariable=self.time_frame_var, width=50)
            self.time_frame_entry.grid(row=8, column=1, padx=5, pady=5)
            self.time_frame_var.trace_add('write', self.update_template_display)  # Trace variable changes

            # New TermLength Field
            self.term_length_var = tk.StringVar()
            tk.Label(frame, text="Term Length:", bg=self.bg_color, fg=self.fg_color).grid(row=9, column=0, sticky=tk.W, padx=5, pady=5)
            self.term_length_entry = tk.Entry(frame, textvariable=self.term_length_var, width=50)
            self.term_length_entry.grid(row=9, column=1, padx=5, pady=5)
            self.term_length_var.trace_add('write', self.update_template_display)  # Trace variable changes

        elif self.prompt_type.get() == "Data Cleansing":
            prompt_text = "(U) DLA Energy Clears Expired Commitments: During the period [Beginning Date] through [End date], Aerospace Energy cleared [total number] expired commitments line items.  As a result, [# of MODs] modifications were processed, with a total value of $[Dollar Value]."
            tk.Label(frame, text="Prompt:", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
            tk.Label(frame, text=prompt_text, wraplength=700, justify="left", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=1, columnspan=2, sticky=tk.W, padx=5, pady=5)

            # Input fields for Data Cleansing prompt
            self.begin_date_var = tk.StringVar()
            tk.Label(frame, text="Beginning Date:", bg=self.bg_color, fg=self.fg_color).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
            self.begin_date_entry = tk.Entry(frame, textvariable=self.begin_date_var, width=50)
            self.begin_date_entry.grid(row=1, column=1, padx=5, pady=5)
            self.begin_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.begin_date_entry))
            self.begin_date_var.trace_add('write', self.update_template_display)

            self.end_date_var = tk.StringVar()
            tk.Label(frame, text="End Date:", bg=self.bg_color, fg=self.fg_color).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
            self.end_date_entry = tk.Entry(frame, textvariable=self.end_date_var, width=50)
            self.end_date_entry.grid(row=2, column=1, padx=5, pady=5)
            self.end_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.end_date_entry))
            self.end_date_var.trace_add('write', self.update_template_display)

            self.total_number_var = tk.StringVar()
            tk.Label(frame, text="Total Number:", bg=self.bg_color, fg=self.fg_color).grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
            self.total_number_entry = tk.Entry(frame, textvariable=self.total_number_var, width=50)
            self.total_number_entry.grid(row=3, column=1, padx=5, pady=5)
            self.total_number_var.trace_add('write', self.update_template_display)

            self.num_mods_var = tk.StringVar()
            tk.Label(frame, text="Number of MODs:", bg=self.bg_color, fg=self.fg_color).grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
            self.num_mods_entry = tk.Entry(frame, textvariable=self.num_mods_var, width=50)
            self.num_mods_entry.grid(row=4, column=1, padx=5, pady=5)
            self.num_mods_var.trace_add('write', self.update_template_display)

            self.dollar_value_var = tk.StringVar()
            tk.Label(frame, text="Dollar Value:", bg=self.bg_color, fg=self.fg_color).grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
            self.dollar_value_entry = tk.Entry(frame, textvariable=self.dollar_value_var, width=50)
            self.dollar_value_entry.grid(row=5, column=1, padx=5, pady=5)
            self.dollar_value_var.trace_add('write', self.update_template_display)

        elif self.prompt_type.get() == "CPARS":
            prompt_text = "(U) DLA Energy Contractor Performance Assessment Reporting System (CPARS) compliance reporting: During the period [Beginning Date], through [End date], Aerospace Energy CPARS compliance percentage is [% of reports]% with [number of reports] reports compliant."
            tk.Label(frame, text="Prompt:", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
            tk.Label(frame, text=prompt_text, wraplength=700, justify="left", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=1, columnspan=2, sticky=tk.W, padx=5, pady=5)

            # Input fields for CPARS prompt
            self.cpars_begin_date_var = tk.StringVar()
            tk.Label(frame, text="Beginning Date:", bg=self.bg_color, fg=self.fg_color).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
            self.cpars_begin_date_entry = tk.Entry(frame, textvariable=self.cpars_begin_date_var, width=50)
            self.cpars_begin_date_entry.grid(row=1, column=1, padx=5, pady=5)
            self.cpars_begin_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.cpars_begin_date_entry))
            self.cpars_begin_date_var.trace_add('write', self.update_template_display)

            self.cpars_end_date_var = tk.StringVar()
            tk.Label(frame, text="End Date:", bg=self.bg_color, fg=self.fg_color).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
            self.cpars_end_date_entry = tk.Entry(frame, textvariable=self.cpars_end_date_var, width=50)
            self.cpars_end_date_entry.grid(row=2, column=1, padx=5, pady=5)
            self.cpars_end_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.cpars_end_date_entry))
            self.cpars_end_date_var.trace_add('write', self.update_template_display)

            self.cpars_percentage_var = tk.StringVar()
            tk.Label(frame, text="% of Reports:", bg=self.bg_color, fg=self.fg_color).grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
            self.cpars_percentage_entry = tk.Entry(frame, textvariable=self.cpars_percentage_var, width=50)
            self.cpars_percentage_entry.grid(row=3, column=1, padx=5, pady=5)
            self.cpars_percentage_var.trace_add('write', self.update_template_display)

            self.cpars_number_reports_var = tk.StringVar()
            tk.Label(frame, text="Number of Reports:", bg=self.bg_color, fg=self.fg_color).grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
            self.cpars_number_reports_entry = tk.Entry(frame, textvariable=self.cpars_number_reports_var, width=50)
            self.cpars_number_reports_entry.grid(row=4, column=1, padx=5, pady=5)
            self.cpars_number_reports_var.trace_add('write', self.update_template_display)

        elif self.prompt_type.get() == "FPDS-NG":
            prompt_text = "(U) DLA Energy Federal Procurement Data System – Next Generation (FPDS-NG) Contract Action Report (CAR) data cleansing: During the period [Beginning Date] through [End date], the Aerospace Energy FPDS-NG Focal Point processed [Number of calls] Call to Action Items (i.e. legacy/current delivery order actions). As a result, [Number of CARs] CARs were finalized in FPDS-NG and Enterprise Business System valued at $[Dollar value]."
            tk.Label(frame, text="Prompt:", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
            tk.Label(frame, text=prompt_text, wraplength=700, justify="left", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=1, columnspan=2, sticky=tk.W, padx=5, pady=5)

            # Input fields for FPDS-NG prompt
            self.fpds_begin_date_var = tk.StringVar()
            tk.Label(frame, text="Beginning Date:", bg=self.bg_color, fg=self.fg_color).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
            self.fpds_begin_date_entry = tk.Entry(frame, textvariable=self.fpds_begin_date_var, width=50)
            self.fpds_begin_date_entry.grid(row=1, column=1, padx=5, pady=5)
            self.fpds_begin_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.fpds_begin_date_entry))
            self.fpds_begin_date_var.trace_add('write', self.update_template_display)

            self.fpds_end_date_var = tk.StringVar()
            tk.Label(frame, text="End Date:", bg=self.bg_color, fg=self.fg_color).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
            self.fpds_end_date_entry = tk.Entry(frame, textvariable=self.fpds_end_date_var, width=50)
            self.fpds_end_date_entry.grid(row=2, column=1, padx=5, pady=5)
            self.fpds_end_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.fpds_end_date_entry))
            self.fpds_end_date_var.trace_add('write', self.update_template_display)

            self.fpds_number_calls_var = tk.StringVar()
            tk.Label(frame, text="Number of Call to Action Items:", bg=self.bg_color, fg=self.fg_color).grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
            self.fpds_number_calls_entry = tk.Entry(frame, textvariable=self.fpds_number_calls_var, width=50)
            self.fpds_number_calls_entry.grid(row=3, column=1, padx=5, pady=5)
            self.fpds_number_calls_var.trace_add('write', self.update_template_display)

            self.fpds_number_cars_var = tk.StringVar()
            tk.Label(frame, text="Number of CARs:", bg=self.bg_color, fg=self.fg_color).grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
            self.fpds_number_cars_entry = tk.Entry(frame, textvariable=self.fpds_number_cars_var, width=50)
            self.fpds_number_cars_entry.grid(row=4, column=1, padx=5, pady=5)
            self.fpds_number_cars_var.trace_add('write', self.update_template_display)

            self.fpds_dollar_value_var = tk.StringVar()
            tk.Label(frame, text="Dollar Value:", bg=self.bg_color, fg=self.fg_color).grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
            self.fpds_dollar_value_entry = tk.Entry(frame, textvariable=self.fpds_dollar_value_var, width=50)
            self.fpds_dollar_value_entry.grid(row=5, column=1, padx=5, pady=5)
            self.fpds_dollar_value_var.trace_add('write', self.update_template_display)

        elif self.prompt_type.get() == "Award":
            prompt_text = "(U) Aerospace Energy Awards Contract for [Product] for [Location]: On [Date], Aerospace Energy awarded contract SPE601-[last 7 contract] for the delivery of [Product] in support of [Location]. The procurement resulted in a [Years], Firm Fixed-Price Requirements-Type contract with a performance period of [p_o_p start], through [p_o_p end date]."
            tk.Label(frame, text="Prompt:", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
            tk.Label(frame, text=prompt_text, wraplength=700, justify="left", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=1, columnspan=2, sticky=tk.W, padx=5, pady=5)

            # Replace Product Entry with Dropdown
            self.award_product_var = tk.StringVar()
            tk.Label(frame, text="Product:", bg=self.bg_color, fg=self.fg_color).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
            self.award_product_dropdown = ttk.Combobox(frame, textvariable=self.award_product_var, values=self.product_list, state="readonly", width=47)
            self.award_product_dropdown.grid(row=1, column=1, padx=5, pady=5)
            self.award_product_var.trace_add('write', self.update_template_display)

            # Consolidated Location Input
            self.award_location_var = tk.StringVar()
            tk.Label(frame, text="Location:", bg=self.bg_color, fg=self.fg_color).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
            self.award_location_entry = tk.Entry(frame, textvariable=self.award_location_var, width=50)
            self.award_location_entry.grid(row=2, column=1, padx=5, pady=5)
            self.award_location_var.trace_add('write', self.update_template_display)

            self.award_date_var = tk.StringVar()
            tk.Label(frame, text="Date:", bg=self.bg_color, fg=self.fg_color).grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
            self.award_date_entry = tk.Entry(frame, textvariable=self.award_date_var, width=50)
            self.award_date_entry.grid(row=3, column=1, padx=5, pady=5)
            self.award_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.award_date_entry))
            self.award_date_var.trace_add('write', self.update_template_display)

            self.award_last_7_contract_var = tk.StringVar()
            tk.Label(frame, text="Last 7 of Contract:", bg=self.bg_color, fg=self.fg_color).grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
            self.award_last_7_contract_entry = tk.Entry(frame, textvariable=self.award_last_7_contract_var, width=50)
            self.award_last_7_contract_entry.grid(row=4, column=1, padx=5, pady=5)
            self.award_last_7_contract_var.trace_add('write', self.update_template_display)

            self.award_years_var = tk.StringVar()
            tk.Label(frame, text="Years:", bg=self.bg_color, fg=self.fg_color).grid(row=6, column=0, sticky=tk.W, padx=5, pady=5)
            self.award_years_entry = tk.Entry(frame, textvariable=self.award_years_var, width=50)
            self.award_years_entry.grid(row=6, column=1, padx=5, pady=5)
            self.award_years_var.trace_add('write', self.update_template_display)

            self.award_p_o_p_start_var = tk.StringVar()
            tk.Label(frame, text="P.O.P Start:", bg=self.bg_color, fg=self.fg_color).grid(row=7, column=0, sticky=tk.W, padx=5, pady=5)
            self.award_p_o_p_start_entry = tk.Entry(frame, textvariable=self.award_p_o_p_start_var, width=50)
            self.award_p_o_p_start_entry.grid(row=7, column=1, padx=5, pady=5)
            self.award_p_o_p_start_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.award_p_o_p_start_entry))
            self.award_p_o_p_start_var.trace_add('write', self.update_template_display)

            self.award_p_o_p_end_date_var = tk.StringVar()
            tk.Label(frame, text="P.O.P End Date:", bg=self.bg_color, fg=self.fg_color).grid(row=8, column=0, sticky=tk.W, padx=5, pady=5)
            self.award_p_o_p_end_date_entry = tk.Entry(frame, textvariable=self.award_p_o_p_end_date_var, width=50)
            self.award_p_o_p_end_date_entry.grid(row=8, column=1, padx=5, pady=5)
            self.award_p_o_p_end_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.award_p_o_p_end_date_entry))
            self.award_p_o_p_end_date_var.trace_add('write', self.update_template_display)

        elif self.prompt_type.get() == "Issues Solicitation / Synopsis/ Sources Sought Notice":
            prompt_text = "(U) Aerospace Energy Issues [Issues Solicitation / Synopsis/ Sources Sought] for [Product] for [Location]: On [Date], Aerospace Energy posted a synopsis on SAM.gov for solicitation [Solicitation number] for the delivery of [Product] including associated services to [Location]. The procurement is anticipated to result a [Term], Firm Fixed-Price Requirements Type contract with a period of performance from [p_o_p Start] through [p_o_p End]. The procurement will be conducted in accordance with lowest price technically acceptable source selection procedures."
            tk.Label(frame, text="Prompt:", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
            tk.Label(frame, text=prompt_text, wraplength=700, justify="left", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=1, columnspan=2, sticky=tk.W, padx=5, pady=5)

            # Input fields for Issues Solicitation prompt
            self.issues_solicitation_type_var = tk.StringVar(value="Issues Solicitation")  # Default Value
            tk.Label(frame, text="Issues Solicitation/Synopsis/Sources Sought:", bg=self.bg_color, fg=self.fg_color).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
            self.issues_solicitation_type_dropdown = ttk.Combobox(frame, textvariable=self.issues_solicitation_type_var, values=["Issues Solicitation", "Synopsis", "Sources Sought Notice"], state="readonly", width=47)
            self.issues_solicitation_type_dropdown.grid(row=1, column=1, padx=5, pady=5)
            self.issues_solicitation_type_var.trace_add('write', self.update_template_display)

            # Replace Product Entry with Dropdown
            self.issues_product_var = tk.StringVar()
            tk.Label(frame, text="Product:", bg=self.bg_color, fg=self.fg_color).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
            self.issues_product_dropdown = ttk.Combobox(frame, textvariable=self.issues_product_var, values=self.product_list, state="readonly", width=47)
            self.issues_product_dropdown.grid(row=2, column=1, padx=5, pady=5)
            self.issues_product_var.trace_add('write', self.update_template_display)

            self.issues_location_var = tk.StringVar()
            tk.Label(frame, text="Location:", bg=self.bg_color, fg=self.fg_color).grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
            self.issues_location_entry = tk.Entry(frame, textvariable=self.issues_location_var, width=50)
            self.issues_location_entry.grid(row=3, column=1, padx=5, pady=5)
            self.issues_location_var.trace_add('write', self.update_template_display)

            self.issues_date_var = tk.StringVar()
            tk.Label(frame, text="Date:", bg=self.bg_color, fg=self.fg_color).grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
            self.issues_date_entry = tk.Entry(frame, textvariable=self.issues_date_var, width=50)
            self.issues_date_entry.grid(row=4, column=1, padx=5, pady=5)
            self.issues_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.issues_date_entry))
            self.issues_date_var.trace_add('write', self.update_template_display)

            self.issues_solicitation_number_var = tk.StringVar()
            tk.Label(frame, text="Solicitation Number:", bg=self.bg_color, fg=self.fg_color).grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
            self.issues_solicitation_number_entry = tk.Entry(frame, textvariable=self.issues_solicitation_number_var, width=50)
            self.issues_solicitation_number_entry.grid(row=5, column=1, padx=5, pady=5)
            self.issues_solicitation_number_var.trace_add('write', self.update_template_display)

            self.issues_term_var = tk.StringVar()
            tk.Label(frame, text="Term:", bg=self.bg_color, fg=self.fg_color).grid(row=6, column=0, sticky=tk.W, padx=5, pady=5)
            self.issues_term_entry = tk.Entry(frame, textvariable=self.issues_term_var, width=50)
            self.issues_term_entry.grid(row=6, column=1, padx=5, pady=5)
            self.issues_term_var.trace_add('write', self.update_template_display)

            self.issues_pop_start_var = tk.StringVar()
            tk.Label(frame, text="P.O.P. Start:", bg=self.bg_color, fg=self.fg_color).grid(row=7, column=0, sticky=tk.W, padx=5, pady=5)
            self.issues_pop_start_entry = tk.Entry(frame, textvariable=self.issues_pop_start_var, width=50)
            self.issues_pop_start_entry.grid(row=7, column=1, padx=5, pady=5)
            self.issues_pop_start_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.issues_pop_start_entry))
            self.issues_pop_start_var.trace_add('write', self.update_template_display)

            self.issues_pop_end_var = tk.StringVar()
            tk.Label(frame, text="P.O.P. End:", bg=self.bg_color, fg=self.fg_color).grid(row=8, column=0, sticky=tk.W, padx=5, pady=5)
            self.issues_pop_end_entry = tk.Entry(frame, textvariable=self.issues_pop_end_var, width=50)
            self.issues_pop_end_entry.grid(row=8, column=1, padx=5, pady=5)
            self.issues_pop_end_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.issues_pop_end_entry))
            self.issues_pop_end_var.trace_add('write', self.update_template_display)

        elif self.prompt_type.get() == "Negotiation":
            prompt_text = "(U) Aerospace Energy [Open/ Close] Negotiations for [Product] for [Location]:  On [Date] , Aerospace Energy [Open/ Close] negotiations for [Solicitation number] delivery of [Product] in support of [Location]. The procurement will result in a [Term], Firm Fixed-Price Requirements-Type contract with a period of performance from [P_O_P start], through [P_O_P End] .  The procurement will be conducted in accordance with the lowest price technically acceptable source selection procedures."
            tk.Label(frame, text="Prompt:", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
            tk.Label(frame, text=prompt_text, wraplength=700, justify="left", bg=self.bg_color, fg=self.fg_color).grid(row=0, column=1, columnspan=2, sticky=tk.W, padx=5, pady=5)

            # Input fields for Negotiation prompt
            self.negotiation_open_close_var = tk.StringVar(value="Open")  # Default value set to "Open"
            tk.Label(frame, text="Open/Close:", bg=self.bg_color, fg=self.fg_color).grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
            self.negotiation_open_close_dropdown = ttk.Combobox(frame, textvariable=self.negotiation_open_close_var, values=["Open", "Close"], state="readonly", width=47) # Adjusted width
            self.negotiation_open_close_dropdown.grid(row=1, column=1, padx=5, pady=5)
            self.negotiation_open_close_var.trace_add('write', self.update_template_display)

            # Replace Product Entry with Dropdown
            self.negotiation_product_var = tk.StringVar()
            tk.Label(frame, text="Product:", bg=self.bg_color, fg=self.fg_color).grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
            self.negotiation_product_dropdown = ttk.Combobox(frame, textvariable=self.negotiation_product_var, values=self.product_list, state="readonly", width=47)
            self.negotiation_product_dropdown.grid(row=2, column=1, padx=5, pady=5)
            self.negotiation_product_var.trace_add('write', self.update_template_display)

            self.negotiation_date_var = tk.StringVar()
            tk.Label(frame, text="Date:", bg=self.bg_color, fg=self.fg_color).grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
            self.negotiation_date_entry = tk.Entry(frame, textvariable=self.negotiation_date_var, width=50)
            self.negotiation_date_entry.grid(row=3, column=1, padx=5, pady=5)
            self.negotiation_date_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.negotiation_date_entry))
            self.negotiation_date_var.trace_add('write', self.update_template_display)

            self.negotiation_solicitation_number_var = tk.StringVar()
            tk.Label(frame, text="Solicitation Number:", bg=self.bg_color, fg=self.fg_color).grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
            self.negotiation_solicitation_number_entry = tk.Entry(frame, textvariable=self.negotiation_solicitation_number_var, width=50)
            self.negotiation_solicitation_number_entry.grid(row=4, column=1, padx=5, pady=5)
            self.negotiation_solicitation_number_var.trace_add('write', self.update_template_display)

            self.negotiation_location_var = tk.StringVar()
            tk.Label(frame, text="Location:", bg=self.bg_color, fg=self.fg_color).grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
            self.negotiation_location_entry = tk.Entry(frame, textvariable=self.negotiation_location_var, width=50)
            self.negotiation_location_entry.grid(row=5, column=1, padx=5, pady=5)
            self.negotiation_location_var.trace_add('write', self.update_template_display)

            self.negotiation_term_var = tk.StringVar()
            tk.Label(frame, text="Term:", bg=self.bg_color, fg=self.fg_color).grid(row=6, column=0, sticky=tk.W, padx=5, pady=5)
            self.negotiation_term_entry = tk.Entry(frame, textvariable=self.negotiation_term_var, width=50)
            self.negotiation_term_entry.grid(row=6, column=1, padx=5, pady=5)
            self.negotiation_term_var.trace_add('write', self.update_template_display)

            self.negotiation_p_o_p_start_var = tk.StringVar()
            tk.Label(frame, text="P.O.P. Start:", bg=self.bg_color, fg=self.fg_color).grid(row=7, column=0, sticky=tk.W, padx=5, pady=5)
            self.negotiation_p_o_p_start_entry = tk.Entry(frame, textvariable=self.negotiation_p_o_p_start_var, width=50)
            self.negotiation_p_o_p_start_entry.grid(row=7, column=1, padx=5, pady=5)
            self.negotiation_p_o_p_start_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.negotiation_p_o_p_start_entry))
            self.negotiation_p_o_p_start_var.trace_add('write', self.update_template_display)

            self.negotiation_p_o_p_end_var = tk.StringVar()
            tk.Label(frame, text="P.O.P. End:", bg=self.bg_color, fg=self.fg_color).grid(row=8, column=0, sticky=tk.W, padx=5, pady=5)
            self.negotiation_p_o_p_end_entry = tk.Entry(frame, textvariable=self.negotiation_p_o_p_end_var, width=50)
            self.negotiation_p_o_p_end_entry.grid(row=8, column=1, padx=5, pady=5)
            self.negotiation_p_o_p_end_entry.bind("<Button-1>", lambda event: self.popup_calendar(self.negotiation_p_o_p_end_entry))
            self.negotiation_p_o_p_end_var.trace_add('write', self.update_template_display)

        else:
            pass #No fields needed for custom.


    def update_prompt_fields(self, event=None):
        # Custom Prompt Section
        if self.prompt_type.get() == "Custom":
            self.custom_prompt_label.grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)  
            self.custom_prompt_entry = tk.Text(self.sitrep_tab, wrap=tk.WORD, width=60, height=3, bg="white", fg="black")
            self.custom_prompt_entry.grid(row=4, column=1, columnspan=2, padx=5, pady=5, sticky=tk.W + tk.E)
            self.custom_prompt_entry.bind("<KeyRelease>", self.update_template_display) # Updates template on key release

            self.save_sitrep_button.grid(row=5, column=1, padx=5, pady=5)  # Below the text input
            self.generate_sitrep_button.grid(row=6, column=1, padx=5, pady=10)  # Below the text input, next to Add Sitrep

            self.create_sitrep_input_fields(self.sitrep_input_frame)
            for widget in self.sitrep_input_frame.winfo_children():
                widget.configure(state='disabled')
            self.sitrep_input_frame.grid_forget()

        else:
            self.custom_prompt_label.grid_forget()
            if hasattr(self, 'custom_prompt_entry'):
                self.custom_prompt_entry.grid_forget()
            self.save_sitrep_button.grid(row=6, column=1, padx=5, pady=5)
            self.generate_sitrep_button.grid(row=7, column=1, padx=5, pady=10)
            self.create_sitrep_input_fields(self.sitrep_input_frame)
            self.sitrep_input_frame.grid(row=5, column=0, columnspan=3, sticky=tk.W + tk.E, padx=5, pady=5)

    def save_sitrep(self):
        # Get current values from input fields
        if self.prompt_type.get() == "Purchase":
            current_data = {
                "PromptType": "Purchase",
                "DoesWhat": self.does_what_var.get().strip(), # <--- ADDED .strip()
                "Item": self.item_var.get().strip(), # <--- ADDED .strip()
                "Place": self.place_var.get().strip(), # <--- ADDED .strip()
                "DateValue": self.date_value_var.get().strip(), # <--- ADDED .strip()
                "DidWhat": self.did_what_var.get().strip(), # <--- ADDED .strip()
                "SpecificItem": self.specific_item_var.get().strip(), # <--- ADDED .strip()
                "CustomerSupportDetail": self.customer_support_detail_var.get().strip(), # <--- ADDED .strip()
                "TimeFrame": self.time_frame_var.get().strip(), # <--- ADDED .strip()
                "TermLength": self.term_length_var.get().strip() # <--- ADDED .strip()
            }
        elif self.prompt_type.get() == "Data Cleansing":
            start_date = self.begin_date_var.get().strip() # <--- ADDED .strip()
            end_date = self.end_date_var.get().strip()     # <--- ADDED .strip()

            if start_date and end_date:
                try:
                    from datetime import datetime
                    start_date_obj = datetime.strptime(start_date, "%B %d, %Y")
                    end_date_obj = datetime.strptime(end_date, "%B %d, %Y")

                    if end_date_obj < start_date_obj:
                        messagebox.showerror("Error", "End Date cannot be before Start Date.")
                        return  # Prevent saving
                except ValueError:
                    messagebox.showerror("Error", "Invalid date format. Please use the calendar.")
                    return

            current_data = {
                "PromptType": "Data Cleansing",
                "Beginning Date": start_date, # Already stripped above
                "End date": end_date,         # Already stripped above
                "total number": self.total_number_var.get().strip(), # <--- ADDED .strip()
                "# of MODs": self.num_mods_var.get().strip(), # <--- ADDED .strip()
                "Dollar Value": self.dollar_value_var.get().strip() # <--- ADDED .strip()
            }

        elif self.prompt_type.get() == "CPARS":
            start_date = self.cpars_begin_date_var.get().strip() # <--- ADDED .strip()
            end_date = self.cpars_end_date_var.get().strip()     # <--- ADDED .strip()

            if start_date and end_date:
                try:
                    from datetime import datetime
                    start_date_obj = datetime.strptime(start_date, "%B %d, %Y")
                    end_date_obj = datetime.strptime(end_date, "%B %d, %Y")

                    if end_date_obj < start_date_obj:
                        messagebox.showerror("Error", "End Date cannot be before Start Date.")
                        return  # Prevent saving
                except ValueError:
                    messagebox.showerror("Error", "Invalid date format. Please use the calendar.")
                    return

            current_data = {
                "PromptType": "CPARS",
                "Beginning Date": start_date, # Already stripped above
                "End date": end_date,         # Already stripped above
                "% of reports": self.cpars_percentage_var.get().strip(), # <--- ADDED .strip()
                "number of reports": self.cpars_number_reports_var.get().strip() # <--- ADDED .strip()
            }

        elif self.prompt_type.get() == "FPDS-NG":
            start_date = self.fpds_begin_date_var.get().strip() # <--- ADDED .strip()
            end_date = self.fpds_end_date_var.get().strip()     # <--- ADDED .strip()

            if start_date and end_date:
                try:
                    from datetime import datetime
                    start_date_obj = datetime.strptime(start_date, "%B %d, %Y")
                    end_date_obj = datetime.strptime(end_date, "%B %d, %Y")

                    if end_date_obj < start_date_obj:
                        messagebox.showerror("Error", "End Date cannot be before Start Date.")
                        return  # Prevent saving
                except ValueError:
                    messagebox.showerror("Error", "Invalid date format. Please use the calendar.")
                    return

            current_data = {
                "PromptType": "FPDS-NG",
                "Beginning Date": start_date, # Already stripped above
                "End date": end_date,         # Already stripped above
                "Number of calls": self.fpds_number_calls_var.get().strip(), # <--- ADDED .strip()
                "Number of CARs": self.fpds_number_cars_var.get().strip(), # <--- ADDED .strip()
                "Dollar value": self.fpds_dollar_value_var.get().strip() # <--- ADDED .strip()
            }

        elif self.prompt_type.get() == "Award":
            # Retrieve P.O.P dates (already stripped in your previous request, keeping it consistent)
            p_o_p_start_date_str = self.award_p_o_p_start_var.get().strip() # <--- ADDED .strip()
            p_o_p_end_date_str = self.award_p_o_p_end_date_var.get().strip() # <--- ADDED .strip()

            # --- Start of new code piece for P.O.P date validation ---
            if p_o_p_start_date_str and p_o_p_end_date_str: # Only validate if both dates are provided
                try:
                    from datetime import datetime # Ensure datetime is available
                    p_o_p_start_date_obj = datetime.strptime(p_o_p_start_date_str, "%B %d, %Y")
                    p_o_p_end_date_obj = datetime.strptime(p_o_p_end_date_str, "%B %d, %Y")

                    if p_o_p_end_date_obj < p_o_p_start_date_obj:
                        messagebox.showerror("Date Error", "P.O.P End Date cannot be before P.O.P Start Date.")
                        return  # Stop the save operation
                except ValueError:
                    # This handles cases where the date string might be malformed, though calendar should prevent it.
                    messagebox.showerror("Date Format Error", "Invalid date format for P.O.P. Please use the calendar to select valid dates.")
                    return # Stop the save operation
            # --- End of new code piece ---

            current_data = {
                "PromptType": "Award",
                "Product": self.award_product_var.get().strip(), # <--- ADDED .strip()
                "Location": self.award_location_var.get().strip(), # <--- ADDED .strip()
                "Date": self.award_date_var.get().strip(), # <--- ADDED .strip()
                "last 7 contract": self.award_last_7_contract_var.get().strip(), # <--- ADDED .strip()
                "Years": self.award_years_var.get().strip(), # <--- ADDED .strip()
                "p_o_p start": p_o_p_start_date_str,  # Already stripped above
                "p_o_p end date": p_o_p_end_date_str    # Already stripped above
            }

        elif self.prompt_type.get() == "Issues Solicitation / Synopsis/ Sources Sought Notice":
            start_date = self.issues_pop_start_var.get().strip() # <--- ADDED .strip()
            end_date = self.issues_pop_end_var.get().strip()     # <--- ADDED .strip()

            if start_date and end_date:
                try:
                    from datetime import datetime
                    start_date_obj = datetime.strptime(start_date, "%B %d, %Y")
                    end_date_obj = datetime.strptime(end_date, "%B %d, %Y")

                    if end_date_obj < start_date_obj:
                        messagebox.showerror("Error", "End Date cannot be before Start Date.")
                        return  # Prevent saving
                except ValueError:
                    messagebox.showerror("Error", "Invalid date format. Please use the calendar.")
                    return

            current_data = {
                "PromptType": "Issues Solicitation / Synopsis/ Sources Sought Notice",
                "Issues Solicitation / Synopsis/ Sources Sought": self.issues_solicitation_type_var.get().strip(), # <--- ADDED .strip()
                "Product": self.issues_product_var.get().strip(), # <--- ADDED .strip()
                "Location": self.issues_location_var.get().strip(), # <--- ADDED .strip()
                "Date": self.issues_date_var.get().strip(), # <--- ADDED .strip()
                "Solicitation number": self.issues_solicitation_number_var.get().strip(), # <--- ADDED .strip()
                "Term": self.issues_term_var.get().strip(), # <--- ADDED .strip()
                "p_o_p Start": start_date, # Already stripped above
                "p_o_p End": end_date      # Already stripped above
            }

        elif self.prompt_type.get() == "Negotiation":
            start_date = self.negotiation_p_o_p_start_var.get().strip() # <--- ADDED .strip()
            end_date = self.negotiation_p_o_p_end_var.get().strip()     # <--- ADDED .strip()

            if start_date and end_date:
                try:
                    from datetime import datetime
                    start_date_obj = datetime.strptime(start_date, "%B %d, %Y")
                    end_date_obj = datetime.strptime(end_date, "%B %d, %Y")

                    if end_date_obj < start_date_obj:
                        messagebox.showerror("Error", "End Date cannot be before Start Date.")
                        return  # Prevent saving
                except ValueError:
                    messagebox.showerror("Error", "Invalid date format. Please use the calendar.")
                    return

            current_data = {
                "PromptType": "Negotiation",
                "Open/ Close": self.negotiation_open_close_var.get().strip(), # <--- ADDED .strip()
                "Product": self.negotiation_product_var.get().strip(), # <--- ADDED .strip()
                "Date": self.negotiation_date_var.get().strip(), # <--- ADDED .strip()
                "Solicitation number": self.negotiation_solicitation_number_var.get().strip(), # <--- ADDED .strip()
                "Location": self.negotiation_location_var.get().strip(), # <--- ADDED .strip()
                "Term": self.negotiation_term_var.get().strip(), # <--- ADDED .strip()
                "P_O_P start": start_date, # Already stripped above
                "P_O_P End": end_date      # Already stripped above
            }

        else: # Custom prompt type
            current_data = {
                "PromptType": "Custom",
                "CustomText": self.custom_prompt_entry.get("1.0", tk.END).strip() # Already has .strip()
            }

        self.sitrep_data.append(current_data)  # Append to the list
        self.clear_input_fields()  # Clear input fields
        self.update_template_display()  # Update display instead of popup


    def clear_input_fields(self):
        """Clears the input fields based on the prompt type."""
        if self.prompt_type.get() == "Purchase":
            self.does_what_var.set("")
            self.item_var.set("")
            self.place_var.set("")
            self.date_value_var.set("")
            self.did_what_var.set("")
            self.specific_item_var.set("")
            self.customer_support_detail_var.set("")
            self.time_frame_var.set("")
            self.term_length_var.set("")
        elif self.prompt_type.get() == "Data Cleansing":
            self.begin_date_var.set("")
            self.end_date_var.set("")
            self.total_number_var.set("")
            self.num_mods_var.set("")
            self.dollar_value_var.set("")
        elif self.prompt_type.get() == "CPARS":
            self.cpars_begin_date_var.set("")
            self.cpars_end_date_var.set("")
            self.cpars_percentage_var.set("")
            self.cpars_number_reports_var.set("")
        elif self.prompt_type.get() == "FPDS-NG":
            self.fpds_begin_date_var.set("")
            self.fpds_end_date_var.set("")
            self.fpds_number_calls_var.set("")
            self.fpds_number_cars_var.set("")
            self.fpds_dollar_value_var.set("")
        elif self.prompt_type.get() == "Award":
            self.award_product_var.set("")
            self.award_location_var.set("")
            self.award_date_var.set("")
            self.award_last_7_contract_var.set("")
            self.award_years_var.set("")
            self.award_p_o_p_start_var.set("")
            self.award_p_o_p_end_date_var.set("")
        elif self.prompt_type.get() == "Issues Solicitation / Synopsis/ Sources Sought Notice":
            self.issues_solicitation_type_var.set("")
            self.issues_product_var.set("")
            self.issues_location_var.set("")
            self.issues_date_var.set("")
            self.issues_solicitation_number_var.set("")
            self.issues_term_var.set("")
            self.issues_pop_start_var.set("")
            self.issues_pop_end_var.set("")
        elif self.prompt_type.get() == "Negotiation":
            self.negotiation_open_close_var.set("")
            self.negotiation_product_var.set("")
            self.negotiation_date_var.set("")
            self.negotiation_solicitation_number_var.set("")
            self.negotiation_location_var.set("")
            self.negotiation_term_var.set("")
            self.negotiation_p_o_p_start_var.set("")
            self.negotiation_p_o_p_end_var.set("")

        else:
            self.custom_prompt_entry.delete("1.0", tk.END)  # Clear the custom text box


    def update_template_display(self, *args):
        self.template_display.config(state=tk.NORMAL) # Enable temporarily to write
        self.template_display.delete("1.0", tk.END)  # Clear existing text
        self.template_display.tag_remove("highlight", "1.0", tk.END)  # Clear previous highlights

        for i, sitrep in enumerate(self.sitrep_data):
            if sitrep["PromptType"] == "Purchase":
                template_text = "(U) Aerospace Energy [{DoesWhat}] for [{Item}] to Support the [{Place}]: On [{DateValue}], Aerospace Energy [{DidWhat}] for solicitation [{SpecificItem}] to support [{Place}] customers with [{CustomerSupportDetail}]. The procurement will result in [{TermLength}], firm fixed-price requirements type contracts with a period of performance from [{TimeFrame}]."

                try:
                    filled_template = template_text.format(**sitrep)
                except KeyError as e:
                    filled_template = f"Error: Missing key {e} in data."

                self.template_display.insert(tk.END, filled_template + "\n\n", f"sitrep_{i}")

            elif sitrep["PromptType"] == "Data Cleansing":
                template_text = "(U) DLA Energy Clears Expired Commitments: During the period [{Beginning Date}] through [{End date}], Aerospace Energy cleared [{total number}] expired commitments line items.  As a result, [{# of MODs}] modifications were processed, with a total value of $[{Dollar Value}]."

                try:
                    filled_template = template_text.format(**sitrep)
                except KeyError as e:
                    filled_template = f"Error: Missing key {e} in data."

                self.template_display.insert(tk.END, filled_template + "\n\n", f"sitrep_{i}")
            elif sitrep["PromptType"] == "CPARS":
                template_text = "(U) DLA Energy Contractor Performance Assessment Reporting System (CPARS) compliance reporting: During the period [{Beginning Date}] through [{End date}], Aerospace Energy CPARS compliance percentage is [{% of reports}]% with [{number of reports}] reports compliant."

                try:
                    filled_template = template_text.format(**sitrep)
                except KeyError as e:
                    filled_template = f"Error: Missing key {e} in data."

                self.template_display.insert(tk.END, filled_template + "\n\n", f"sitrep_{i}")
            elif sitrep["PromptType"] == "FPDS-NG":
                template_text = "(U) DLA Energy Federal Procurement Data System – Next Generation (FPDS-NG) Contract Action Report (CAR) data cleansing: During the period [{Beginning Date}] through [{End date}], the Aerospace Energy FPDS-NG Focal Point processed [{Number of calls}] Call to Action Items (i.e. legacy/current delivery order actions). As a result, [{Number of CARs}] CARs were finalized in FPDS-NG and Enterprise Business System valued at $[{Dollar value}]."

                try:
                    filled_template = template_text.format(**sitrep)
                except KeyError as e:
                    filled_template = f"Error: Missing key {e} in data."

                self.template_display.insert(tk.END, filled_template + "\n\n", f"sitrep_{i}")
            elif sitrep["PromptType"] == "Award":
                template_text = "(U) Aerospace Energy Awards Contract for [{Product}] for [{Location}]: On [{Date}], Aerospace Energy awarded contract SPE601-[{last 7 contract}] for the delivery of [{Product}] in support of [{Location}]. The procurement resulted in a [{Years}], Firm Fixed-Price Requirements-Type contract with a performance period of [{p_o_p start}], through [{p_o_p end date}]."

                try:
                    # Use the consolidated variables for both locations
                    filled_template = template_text.format(
                        Product=sitrep.get("Product", ""),  # Use .get() for safety
                        Location=sitrep.get("Location", ""),
                        Date=sitrep.get("Date", ""),
                        **{k: v for k, v in sitrep.items() if k not in ("Product", "Location", "Date")}
                    )
                except KeyError as e:
                    filled_template = f"Error: Missing key {e} in data."

                self.template_display.insert(tk.END, filled_template + "\n\n", f"sitrep_{i}")

            elif sitrep["PromptType"] == "Issues Solicitation / Synopsis/ Sources Sought Notice":
                template_text = "(U) Aerospace Energy Issues [{Issues Solicitation / Synopsis/ Sources Sought}] for [{Product}] for [{Location}]: On [{Date}], Aerospace Energy posted a synopsis on SAM.gov for solicitation [{Solicitation number}] for the delivery of [{Product}] including associated services to [{Location}]. The procurement is anticipated to result a [{Term}], Firm Fixed-Price Requirements Type contract with a period of performance from [{p_o_p Start}] through [{p_o_p End}]. The procurement will be conducted in accordance with lowest price technically acceptable source selection procedures."

                try:
                    filled_template = template_text.format(**sitrep)
                except KeyError as e:
                    filled_template = f"Error: Missing key {e} in data."

                self.template_display.insert(tk.END, filled_template + "\n\n", f"sitrep_{i}")


            elif sitrep["PromptType"] == "Negotiation":
                template_text = "(U) Aerospace Energy [{Open/ Close}] Negotiations for [{Product}] for [{Location}]:  On [{Date}], Aerospace Energy [{Open/ Close}] negotiations for [{Solicitation number}] delivery of [{Product}] in support of [{Location}]. The procurement will result in a [{Term}], Firm Fixed-Price Requirements-Type contract with a period of performance from [{P_O_P start}], through [{P_O_P End}]. The procurement will be conducted in accordance with the lowest price technically acceptable source selection procedures."

                try:
                    filled_template = template_text.format(**sitrep)
                except KeyError as e:
                    filled_template = f"Error: Missing key {e} in data."

                self.template_display.insert(tk.END, filled_template + "\n\n", f"sitrep_{i}")


            else:
                custom_text = sitrep["CustomText"]
                parts = custom_text.split(":", 1)
                bold_part = parts[0] + ":"  # Ensure the colon is included in the bold part
                rest_of_template = parts[1] if len(parts) > 1 else ""
                display_text = f"{bold_part}{rest_of_template}" # Markdown for bold
                self.template_display.insert(tk.END, display_text + "\n\n", f"sitrep_{i}")

        self.template_display.config(state=tk.DISABLED) # Disable again



    def create_word_document(self, output_file):
        """Generates a Word document with the formatted SITREP data."""

        document = docx.Document()

        # Set default font to Times New Roman, Size 12 in the word document
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Loop through each SITREP and add it to the document
        for sitrep in self.sitrep_data:
            #print(f"Processing SITREP: {sitrep}")  # Debug: Print each SITREP values in the terminal can see each input the user inputs and where the value is stored. Looped for each Sitrep created
            try:  # Wrap the entire SITREP processing in a try block
                if sitrep["PromptType"] == "Purchase":
                    template_text = "(U) Aerospace Energy {DoesWhat} for {Item} to Support the {Place}: On {DateValue}, Aerospace Energy {DidWhat} for solicitation {SpecificItem} to support {Place} customers with {CustomerSupportDetail}. The procurement will result in {TermLength}, firm fixed-price requirements type contracts with a period of performance from {TimeFrame}."                    
                    filled_template = template_text.format(**sitrep)
                elif sitrep["PromptType"] == "Data Cleansing":
                    template_text = "(U) DLA Energy Clears Expired Commitments: During the period {Beginning Date} through {End date}, Aerospace Energy cleared {total number} expired commitments line items.  As a result, {# of MODs} modifications were processed, with a total value of ${Dollar Value}."
                    filled_template = template_text.format(**sitrep)
                elif sitrep["PromptType"] == "CPARS":
                    template_text = "(U) DLA Energy Contractor Performance Assessment Reporting System (CPARS) compliance reporting: During the period {Beginning Date} through {End date}, Aerospace Energy CPARS compliance percentage is {% of reports}% with {number of reports} reports compliant."
                    filled_template = template_text.format(**sitrep)
                elif sitrep["PromptType"] == "FPDS-NG":
                    template_text = "(U) DLA Energy Federal Procurement Data System – Next Generation (FPDS-NG) Contract Action Report (CAR) data cleansing: During the period {Beginning Date} through {End date}, the Aerospace Energy FPDS-NG Focal Point processed {Number of calls} Call to Action Items (i.e. legacy/current delivery order actions). As a result, {Number of CARs} CARs were finalized in FPDS-NG and Enterprise Business System valued at ${Dollar value}."
                    filled_template = template_text.format(**sitrep)
                elif sitrep["PromptType"] == "Award":
                    template_text = "(U) Aerospace Energy Awards Contract for {Product} for {Location}: On {Date}, Aerospace Energy awarded contract SPE601-{last 7 contract} for the delivery of {Product} in support of {Location}. The procurement resulted in a {Years}, Firm Fixed-Price Requirements-Type contract with a performance period of {p_o_p start}, through {p_o_p end date}."
                    # Use the consolidated product and location variables
                    filled_template = template_text.format(
                        Product=sitrep.get("Product", ""),
                        Location=sitrep.get("Location", ""),
                        Date=sitrep.get("Date", ""),
                        **{k: v for k, v in sitrep.items() if k not in ("Product", "Location", "Date")}
                    )

                elif sitrep["PromptType"] == "Issues Solicitation / Synopsis/ Sources Sought Notice":
                    template_text = "(U) Aerospace Energy Issues {Issues Solicitation / Synopsis/ Sources Sought} for {Product} for {Location}: On {Date}, Aerospace Energy posted a synopsis on SAM.gov for solicitation {Solicitation number} for the delivery of {Product} including associated services to {Location}. The procurement is anticipated to result a {Term}, Firm Fixed-Price Requirements Type contract with a period of performance from {p_o_p Start} through {p_o_p End}. The procurement will be conducted in accordance with lowest price technically acceptable source selection procedures."
                    filled_template = template_text.format(**sitrep)

                elif sitrep["PromptType"] == "Negotiation":
                    template_text = "(U) Aerospace Energy {Open/ Close} Negotiations for {Product} for {Location}:  On {Date}, Aerospace Energy {Open/ Close} negotiations for {Solicitation number} delivery of {Product} in support of {Location}. The procurement will result in a {Term}, Firm Fixed-Price Requirements-Type contract with a period of performance from {P_O_P start}, through {P_O_P End}. The procurement will be conducted in accordance with the lowest price technically acceptable source selection procedures."
                    filled_template = template_text.format(**sitrep)

                else:
                    filled_template = sitrep["CustomText"]

                # Split the template to bold the first part
                parts = filled_template.split(":", 1)
                bold_part = parts[0] + ":"  # Ensure the colon is included in the bold part
                rest_of_template = parts[1] if len(parts) > 1 else ""

                paragraph = document.add_paragraph()

                # Add the bold part
                run = paragraph.add_run(bold_part)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

                # Add the rest of the template
                run = paragraph.add_run(rest_of_template)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Align left
                paragraph.paragraph_format.space_after = Pt(12)  # Add space after paragraph

            except KeyError as e:
                print(f"KeyError: {e} in SITREP {sitrep}")  # More specific error message
                messagebox.showerror("Error", f"KeyError: {e} in SITREP data.  Check your input fields.") #display error on ui

            except Exception as e:
                print(f"General Error processing SITREP: {e} - SITREP Data: {sitrep}")
                messagebox.showerror("Error", f"An error occurred processing a SITREP: {e}")


        # Save the document
        document.save(output_file)


    def select_sitrep(self, event):
        """Highlights the entire SITREP paragraph that was clicked and loads it for editing."""
        if self.template_display["state"] == tk.DISABLED:
            self.template_display.config(state=tk.NORMAL)  # Temporarily enable

        # Remove any previous highlighting
        self.template_display.tag_remove("highlight", "1.0", tk.END)

        # Get the index of the SITREP clicked
        index = None
        for i, sitrep in enumerate(self.sitrep_data):
            tag_name = f"sitrep_{i}"
            if self.template_display.tag_ranges(tag_name): #check if this tag exists
                start = self.template_display.tag_ranges(tag_name)[0]
                end = self.template_display.tag_ranges(tag_name)[1]
                if self.template_display.index(tk.CURRENT) >= self.template_display.index(start) and self.template_display.index(tk.CURRENT) <= self.template_display.index(end):
                    index = i
                    break

        if index is not None:
            self.selected_sitrep_index = index
            tag_name = f"sitrep_{index}"
            start = self.template_display.tag_ranges(tag_name)[0]
            end = self.template_display.tag_ranges(tag_name)[1]

            # Apply highlighting
            self.template_display.tag_add("highlight", start, end)
            self.template_display.tag_config("highlight", background="lightblue")  # Choose a highlight color

            # Load the selected SITREP into the input fields for editing
            self.load_sitrep_for_editing(self.selected_sitrep_index)

        self.template_display.config(state=tk.DISABLED)  # Disable again

    def load_sitrep_for_editing(self, index):
        """Loads a saved SITREP into the input fields for editing."""
        sitrep = self.sitrep_data[index]

        if sitrep["PromptType"] == "Purchase":
            self.prompt_type.set("Purchase")
            self.does_what_var.set(sitrep.get("DoesWhat", ""))
            # Set item dropdown value
            self.item_var.set(sitrep.get("Item", ""))
            self.place_var.set(sitrep.get("Place", ""))
            self.date_value_var.set(sitrep.get("DateValue", ""))
            self.did_what_var.set(sitrep.get("DidWhat", ""))
            self.specific_item_var.set(sitrep.get("SpecificItem", ""))
            self.customer_support_detail_var.set(sitrep.get("CustomerSupportDetail", ""))
            self.time_frame_var.set(sitrep.get("TimeFrame", ""))
            self.term_length_var.set(sitrep.get("TermLength", ""))


        elif sitrep["PromptType"] == "Data Cleansing":
            self.prompt_type.set("Data Cleansing")
            self.begin_date_var.set(sitrep.get("Beginning Date", ""))
            self.end_date_var.set(sitrep.get("End date", ""))
            self.total_number_var.set(sitrep.get("total number", ""))
            self.num_mods_var.set(sitrep.get("# of MODs", ""))
            self.dollar_value_var.set(sitrep.get("Dollar Value", ""))
        elif sitrep["PromptType"] == "CPARS":
            self.prompt_type.set("CPARS")
            self.cpars_begin_date_var.set(sitrep.get("Beginning Date", ""))
            self.cpars_end_date_var.set(sitrep.get("End date", ""))
            self.cpars_percentage_var.set(sitrep.get("% of reports", ""))
            self.cpars_number_reports_var.set(sitrep.get("number of reports", ""))
        elif sitrep["PromptType"] == "FPDS-NG":
            self.prompt_type.set("FPDS-NG")
            self.fpds_begin_date_var.set(sitrep.get("Beginning Date", ""))
            self.fpds_end_date_var.set(sitrep.get("End date", ""))
            self.fpds_number_calls_var.set(sitrep.get("Number of calls", ""))
            self.fpds_number_cars_var.set(sitrep.get("Number of CARs", ""))
            self.fpds_dollar_value_var.set(sitrep.get("Dollar value", ""))
        elif sitrep["PromptType"] == "Award":
            self.prompt_type.set("Award")
            # Set product dropdown value
            self.award_product_var.set(sitrep.get("Product", ""))
            self.award_location_var.set(sitrep.get("Location", ""))
            self.award_date_var.set(sitrep.get("Date", ""))
            self.award_last_7_contract_var.set(sitrep.get("last 7 contract", ""))
            self.award_years_var.set(sitrep.get("Years", ""))
            self.award_pop_start_var.set(sitrep.get("P_O_P_start", ""))
            self.award_pop_end_date_var.set(sitrep.get("P_O_P_end_date", ""))

        elif sitrep["PromptType"] == "Issues Solicitation / Synopsis/ Sources Sought Notice":
            self.prompt_type.set("Issues Solicitation / Synopsis/ Sources Sought Notice")
            self.issues_solicitation_type_var.set(sitrep.get("Issues Solicitation / Synopsis/ Sources Sought", ""))
            #set product dropdown value
            self.issues_product_var.set(sitrep.get("Product", ""))
            self.issues_location_var.set(sitrep.get("Location", ""))
            self.issues_date_var.set(sitrep.get("Date", ""))
            self.issues_solicitation_number_var.set(sitrep.get("Solicitation number", ""))
            self.issues_term_var.set(sitrep.get("Term", ""))
            self.issues_pop_start_var.set(sitrep.get("P_O_P_Start", ""))
            self.issues_pop_end_var.set(sitrep.get("P_O_P_End", ""))

        elif sitrep["PromptType"] == "Negotiation":
            self.prompt_type.set("Negotiation")
            self.negotiation_open_close_var.set(sitrep.get("Open/ Close", ""))
            # Set product dropdown value
            self.negotiation_product_var.set(sitrep.get("Product", ""))
            self.negotiation_date_var.set(sitrep.get("Date", ""))
            self.negotiation_solicitation_number_var.set(sitrep.get("Solicitation number", ""))
            self.negotiation_location_var.set(sitrep.get("Location", ""))
            self.negotiation_term_var.set(sitrep.get("Term", ""))
            self.negotiation_p_o_p_start_var.set(sitrep.get("P_O_P start", ""))
            self.negotiation_p_o_p_end_var.set(sitrep.get("P_O_P End", ""))


        else:
            self.prompt_type.set("Custom")
            self.custom_prompt_entry.delete("1.0", tk.END)
            self.custom_prompt_entry.insert("1.0", sitrep.get("CustomText", ""))

        self.update_prompt_fields()  # Update UI based on the loaded prompt type

    def delete_selected_sitrep(self):
        """Deletes the currently selected SITREP."""
        if self.selected_sitrep_index is not None:
            del self.sitrep_data[self.selected_sitrep_index]
            self.update_template_display()
            self.selected_sitrep_index = None  # Reset selection
            messagebox.showinfo("Success", "SITREP deleted successfully!")
        else:
            messagebox.showerror("Error", "Select a SITREP to delete.")

    def generate_sitrep(self):
        output_folder = self.output_folder_path.get()
        output_filename = self.output_filename.get()

        if not output_folder:
            messagebox.showerror("Error", "Please select an output folder.")
            return
        if not output_filename:
            messagebox.showerror("Error", "Please enter an output filename.")
            return

        if not output_filename.endswith(".docx"):
            output_filename += ".docx"

        output_file = os.path.join(output_folder, output_filename)

        if not self.sitrep_data:
            messagebox.showerror("Error", "Please add at least one SITREP.")
            return

        try:
            self.create_word_document(output_file)  # Generate the Word document

            messagebox.showinfo("Success", f"SITREP data added successfully!\nSaved to: {output_file}")

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

        self.update_template_display()  # Update display after generating
        self.sitrep_data = []  # Clear data after generating
        self.selected_sitrep_index = None # Reset selected index

root = tk.Tk()
my_gui = SitrepGenerator(root)
root.mainloop()

